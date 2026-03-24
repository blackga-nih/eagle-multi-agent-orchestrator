"""Extract metadata from S3 template files (DOCX/XLSX/PDF/DOC/TXT).

Downloads all templates from the eagle-knowledge-base S3 bucket,
parses each file to extract section structure, placeholders, and table info,
then outputs JSON metadata files to eagle-plugin/data/template-metadata/.

Usage:
    cd server && python ../scripts/extract_template_metadata.py --profile eagle
"""
from __future__ import annotations

import argparse
import io
import json
import logging
import os
import re
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Optional

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── Constants ──
BUCKET = os.getenv("TEMPLATE_BUCKET", "eagle-documents-695681773636-dev")
PREFIX = "eagle-knowledge-base/approved/supervisor-core/essential-templates"

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "eagle-plugin" / "data" / "template-metadata"

# Category inference from filename patterns
CATEGORY_PATTERNS: list[tuple[str, str, str]] = [
    # (regex pattern on filename, category, variant)
    (r"(?i)statement.of.work|sow", "sow", ""),
    (r"(?i)acquisition.plan|^AP[\s_]|S-AP", "acquisition_plan", ""),
    (r"(?i)task.order.acquisition", "acquisition_plan", "task_order"),
    (r"(?i)AP.Under.SAT|1\.a\.", "acquisition_plan", "under_sat"),
    (r"(?i)AP.Above.SAT|1\.b", "acquisition_plan", "above_sat"),
    (r"(?i)streamlined.acquisition.plan|streamlined.ap", "acquisition_plan", "streamlined"),
    (r"(?i)Attch.*HHS.*Streamlined.*AP", "acquisition_plan", "hhs_streamlined"),
    (r"(?i)SON.*Products|3\.a\.", "son_products", ""),
    (r"(?i)SON.*Services|3\.b\.", "son_services", ""),
    (r"(?i)IGE.*Products|4\.a\.", "igce", "products"),
    (r"(?i)IGE.*Services|4\.b\.", "igce", "services"),
    (r"(?i)IGCE|IGE.*Commercial", "igce", ""),
    (r"(?i)Single.Source.J&A|6\.a\.", "justification", "single_source_under_sat"),
    (r"(?i)Justification.*Approval|J&A|J_and_A", "justification", ""),
    (r"(?i)Market.Research|FY2\d.*MR", "market_research", ""),
    (r"(?i)Buy.American.*Non.Avail", "buy_american", "non_availability"),
    (r"(?i)Buy.American.*Other", "buy_american", "other_exceptions"),
    (r"(?i)Conference.*Request$|Attachment.A.*Conference", "conference_request", ""),
    (r"(?i)Conference.*Waiver|Attachment.B.*Conference", "conference_waiver", ""),
    (r"(?i)Promotional.Item|Attachment.D", "promotional_item", ""),
    (r"(?i)Exemption.Determination|Attachment.G", "exemption_determination", ""),
    (r"(?i)Mandatory.Use.Waiver", "mandatory_use_waiver", ""),
    (r"(?i)GFP.Form", "gfp_form", ""),
    (r"(?i)LSJ.*GSA.*BPA|Call.?Orders", "bpa_call_order", ""),
    (r"(?i)Quotation.Abstract", "quotation_abstract", ""),
    (r"(?i)Receiving.Report", "receiving_report", ""),
    (r"(?i)SRB.Request", "srb_request", ""),
    (r"(?i)subk.*review", "subk_review", ""),
    (r"(?i)SubK.*Plan", "subk_plan", ""),
    (r"(?i)Technical.Questionnaire|Project.Officer", "technical_questionnaire", ""),
    (r"(?i)COR.*Appointment|COR.*Memorandum", "cor_certification", ""),
    (r"(?i)AP.*Structure.*Guide", "reference_guide", "ap_structure"),
    (r"(?i)Streamlined.*MR.*Template.*txt", "reference_guide", "mr_template"),
]


@dataclass
class SectionMeta:
    number: str
    title: str
    has_table: bool = False
    placeholders: list[str] = field(default_factory=list)


@dataclass
class TemplateMeta:
    filename: str
    format: str
    category: str
    variant: str
    sections: list[dict] = field(default_factory=list)
    total_placeholders: int = 0
    total_sections: int = 0
    sheet_names: list[str] = field(default_factory=list)
    parse_error: Optional[str] = None


def infer_category(filename: str) -> tuple[str, str]:
    """Infer category and variant from filename."""
    for pattern, category, variant in CATEGORY_PATTERNS:
        if re.search(pattern, filename):
            return category, variant
    return "unknown", ""


def extract_placeholders(text: str) -> list[str]:
    """Extract {{PLACEHOLDER}} tokens from text."""
    return list(dict.fromkeys(re.findall(r"\{\{(\w+)\}\}", text)))


def has_table_markers(text: str) -> bool:
    """Check if text contains markdown table delimiters or Word-style table hints."""
    return bool(re.search(r"\|.*\|.*\|", text))


# ── DOCX Parser ──
def parse_docx(data: bytes, filename: str) -> TemplateMeta:
    """Extract metadata from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return TemplateMeta(
            filename=filename, format="docx",
            category="unknown", variant="",
            parse_error="python-docx not installed",
        )

    category, variant = infer_category(filename)
    doc = Document(io.BytesIO(data))

    sections: list[SectionMeta] = []
    current_section: Optional[SectionMeta] = None
    all_placeholders: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Detect headings
        is_heading = "Heading" in style_name
        if not is_heading:
            # Also detect bold-only short lines as pseudo-headings
            heading_match = re.match(
                r"^(?:(?:PART\s+)?(\d+(?:\.\d+)*)[\s.:)\-]+)?\s*([A-Z][A-Z\s,/&()\-]+)$",
                text,
            )
            if heading_match and len(text) < 120:
                is_heading = True

        if is_heading:
            # Try to extract section number and title
            m = re.match(
                r"^(?:(?:PART\s+|Section\s+)?(\d+(?:\.\d+)*)[\s.:)\-]+)?\s*(.+)$",
                text, re.IGNORECASE,
            )
            if m:
                num = m.group(1) or str(len(sections) + 1)
                title = m.group(2).strip()
            else:
                num = str(len(sections) + 1)
                title = text

            current_section = SectionMeta(number=num, title=title)
            sections.append(current_section)
        else:
            # Check for placeholders and tables in body text
            phs = extract_placeholders(text)
            all_placeholders.extend(phs)
            if current_section:
                current_section.placeholders.extend(phs)
                if has_table_markers(text):
                    current_section.has_table = True

    # Also scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                phs = extract_placeholders(cell.text)
                all_placeholders.extend(phs)
                if current_section:
                    current_section.placeholders.extend(phs)
                    current_section.has_table = True
                elif sections:
                    sections[-1].placeholders.extend(phs)
                    sections[-1].has_table = True

    # Deduplicate placeholders per section
    for s in sections:
        s.placeholders = list(dict.fromkeys(s.placeholders))

    unique_placeholders = list(dict.fromkeys(all_placeholders))

    return TemplateMeta(
        filename=filename,
        format="docx",
        category=category,
        variant=variant,
        sections=[asdict(s) for s in sections],
        total_placeholders=len(unique_placeholders),
        total_sections=len(sections),
    )


# ── XLSX Parser ──
def parse_xlsx(data: bytes, filename: str) -> TemplateMeta:
    """Extract metadata from an XLSX file."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return TemplateMeta(
            filename=filename, format="xlsx",
            category="unknown", variant="",
            parse_error="openpyxl not installed",
        )

    category, variant = infer_category(filename)
    wb = load_workbook(io.BytesIO(data), data_only=True)

    sections: list[SectionMeta] = []
    all_placeholders: list[str] = []
    sheet_names: list[str] = []

    for sheet in wb.worksheets:
        sheet_names.append(sheet.title)
        sheet_section = SectionMeta(
            number=str(len(sections) + 1),
            title=sheet.title,
            has_table=True,  # Spreadsheets are inherently tabular
        )

        for row in sheet.iter_rows(max_row=200):
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    phs = extract_placeholders(cell.value)
                    all_placeholders.extend(phs)
                    sheet_section.placeholders.extend(phs)

        sheet_section.placeholders = list(dict.fromkeys(sheet_section.placeholders))
        sections.append(sheet_section)

    unique_placeholders = list(dict.fromkeys(all_placeholders))

    return TemplateMeta(
        filename=filename,
        format="xlsx",
        category=category,
        variant=variant,
        sections=[asdict(s) for s in sections],
        total_placeholders=len(unique_placeholders),
        total_sections=len(sections),
        sheet_names=sheet_names,
    )


# ── PDF Parser ──
def parse_pdf(data: bytes, filename: str) -> TemplateMeta:
    """Extract metadata from a PDF file."""
    category, variant = infer_category(filename)

    try:
        import pdfplumber
        pdf = pdfplumber.open(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        pdf.close()
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            return TemplateMeta(
                filename=filename, format="pdf",
                category=category, variant=variant,
                parse_error="Neither pdfplumber nor PyPDF2 installed",
            )

    return _parse_plain_text(text, filename, "pdf", category, variant)


# ── TXT Parser ──
def parse_txt(data: bytes, filename: str) -> TemplateMeta:
    """Extract metadata from a plain text file."""
    category, variant = infer_category(filename)
    text = data.decode("utf-8", errors="replace")
    return _parse_plain_text(text, filename, "txt", category, variant)


def _parse_plain_text(
    text: str, filename: str, fmt: str, category: str, variant: str,
) -> TemplateMeta:
    """Parse plain text into sections by detecting heading-like lines."""
    sections: list[SectionMeta] = []
    all_placeholders: list[str] = []
    current_section: Optional[SectionMeta] = None

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Heading detection: numbered sections or ALL CAPS short lines
        heading_match = re.match(
            r"^(?:(?:PART\s+|Section\s+)?(\d+(?:\.\d+)*)[\s.:)\-]+)?\s*([A-Z][A-Z\s,/&()\-]{4,})$",
            stripped,
        )
        if heading_match and len(stripped) < 120:
            num = heading_match.group(1) or str(len(sections) + 1)
            title = heading_match.group(2).strip()
            current_section = SectionMeta(number=num, title=title)
            sections.append(current_section)
        else:
            phs = extract_placeholders(stripped)
            all_placeholders.extend(phs)
            if current_section:
                current_section.placeholders.extend(phs)
                if has_table_markers(stripped):
                    current_section.has_table = True

    for s in sections:
        s.placeholders = list(dict.fromkeys(s.placeholders))

    unique_placeholders = list(dict.fromkeys(all_placeholders))

    return TemplateMeta(
        filename=filename,
        format=fmt,
        category=category,
        variant=variant,
        sections=[asdict(s) for s in sections],
        total_placeholders=len(unique_placeholders),
        total_sections=len(sections),
    )


# ── DOC Parser (legacy .doc via textract or antiword fallback) ──
def parse_doc(data: bytes, filename: str) -> TemplateMeta:
    """Extract metadata from a legacy .doc file."""
    category, variant = infer_category(filename)

    # Try antiword first, then textract
    import subprocess
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    text = ""
    try:
        result = subprocess.run(
            ["antiword", tmp_path], capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0:
            text = result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    if not text:
        try:
            import textract
            text = textract.process(tmp_path).decode("utf-8", errors="replace")
        except ImportError:
            pass

    os.unlink(tmp_path)

    if not text:
        return TemplateMeta(
            filename=filename, format="doc",
            category=category, variant=variant,
            parse_error="Could not parse .doc (antiword/textract unavailable)",
        )

    return _parse_plain_text(text, filename, "doc", category, variant)


# ── Dispatcher ──
PARSERS = {
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".pdf": parse_pdf,
    ".txt": parse_txt,
    ".doc": parse_doc,
}


def parse_template(data: bytes, filename: str) -> TemplateMeta:
    """Route to the correct parser based on file extension."""
    ext = os.path.splitext(filename)[1].lower()
    parser = PARSERS.get(ext)
    if not parser:
        category, variant = infer_category(filename)
        return TemplateMeta(
            filename=filename, format=ext.lstrip("."),
            category=category, variant=variant,
            parse_error=f"No parser for {ext} files",
        )
    return parser(data, filename)


def main():
    parser = argparse.ArgumentParser(description="Extract S3 template metadata")
    parser.add_argument("--profile", default=None, help="AWS profile name")
    parser.add_argument("--bucket", default=BUCKET, help="S3 bucket")
    parser.add_argument("--prefix", default=PREFIX, help="S3 prefix")
    parser.add_argument("--output-dir", default=str(OUTPUT_DIR), help="Output directory")
    parser.add_argument("--dry-run", action="store_true", help="List files without downloading")
    args = parser.parse_args()

    import boto3

    session_kwargs: dict[str, Any] = {}
    if args.profile:
        session_kwargs["profile_name"] = args.profile
    session = boto3.Session(**session_kwargs)
    s3 = session.client("s3", region_name="us-east-1")

    # List all objects under the prefix
    logger.info("Listing s3://%s/%s ...", args.bucket, args.prefix)
    paginator = s3.get_paginator("list_objects_v2")
    files: list[str] = []
    for page in paginator.paginate(Bucket=args.bucket, Prefix=args.prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            filename = os.path.basename(key)
            if filename and not filename.startswith("."):
                files.append(key)

    logger.info("Found %d template files", len(files))

    if args.dry_run:
        for f in files:
            print(f"  {os.path.basename(f)}")
        return

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    results: list[dict] = []
    for s3_key in files:
        filename = os.path.basename(s3_key)
        logger.info("Processing: %s", filename)

        try:
            response = s3.get_object(Bucket=args.bucket, Key=s3_key)
            data = response["Body"].read()
        except Exception as e:
            logger.error("Failed to download %s: %s", filename, e)
            results.append(asdict(TemplateMeta(
                filename=filename, format="unknown",
                category="unknown", variant="",
                parse_error=str(e),
            )))
            continue

        meta = parse_template(data, filename)
        meta_dict = asdict(meta)
        results.append(meta_dict)

        # Write individual JSON
        safe_name = re.sub(r"[^\w\-.]", "_", os.path.splitext(filename)[0])
        out_path = output_dir / f"{safe_name}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(meta_dict, f, indent=2)
        logger.info("  → %s (%d sections, %d placeholders)",
                     out_path.name, meta.total_sections, meta.total_placeholders)

    # Write combined index
    index_path = output_dir / "_index.json"
    index_data = {
        "total_templates": len(results),
        "by_category": {},
        "templates": results,
    }
    for r in results:
        cat = r["category"]
        if cat not in index_data["by_category"]:
            index_data["by_category"][cat] = []
        index_data["by_category"][cat].append(r["filename"])

    with open(index_path, "w", encoding="utf-8") as f:
        json.dump(index_data, f, indent=2)

    logger.info("\nDone! Wrote %d metadata files + _index.json to %s", len(results), output_dir)

    # Summary
    categories = {}
    for r in results:
        cat = r["category"]
        categories[cat] = categories.get(cat, 0) + 1
    logger.info("Categories: %s", json.dumps(categories, indent=2))


if __name__ == "__main__":
    main()
