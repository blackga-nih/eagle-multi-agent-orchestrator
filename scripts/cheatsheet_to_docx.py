"""Convert the 2026-04-23 action-items cheat sheet from markdown to docx.

Kept intentionally narrow — this is not a general markdown→docx converter.
It renders the specific structure of 20260423-action-items-cheatsheet-v1.md:
one legend table, one items table, a roll-up paragraph, and a numbered
next-steps list. Re-run after editing the .md to refresh the .docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(
    "docs/development/meeting-transcripts/20260416-eagle-output-review"
)
DEST = ROOT / "20260423-action-items-cheatsheet-v1.docx"

STATUS_LEGEND = [
    ("✅ Validated", "Observed working on the deployed system or in source of truth"),
    ("🛠 Patched", "Code shipped, behavior not yet verified in use"),
    ("⛔ Outstanding", "No work yet"),
    ("🔄 Ongoing", "Stylistic guardrail, not a one-time task"),
]

ITEMS = [
    ("1",  "KB sync — RFO only, legacy FAR stripped", "✅",
     "S3 bucket verified 2026-04-22: eagle-documents-695681773636-dev/eagle-knowledge-base/approved/ is RFO-only"),
    ("1b", "Q4/Q5 regression root-cause (KB correct but still wrong answer)", "⛔",
     "Need to re-run against current KB and trace why"),
    ("2",  "IGCE/budget reconcile bug", "🛠",
     "PR #150 — matrix.budget_semantics + supervisor forbids reconcile Q verbatim"),
    ("3",  "FFP hybrid + task-area decomposition", "🛠",
     "PR #149 — NIH/NCI T&M institutional override in supervisor; task-area decomp portion still pending"),
    ("4",  "Remove Labor Hour default", "✅",
     'PR #149 — supervisor prompt: "do NOT default to Labor-Hour" (visible in source)'),
    ("5",  "IGCE methodology / budget narrative", "🛠",
     "PR #149 — template gained §2.4 Rate Derivation + §2.5 Budget Narrative; needs observed output to validate"),
    ("6",  "Rewrite Q1–Q5 in natural CO voice", "⛔",
     "Pending Ingrid/Ryan working session"),
    ("7",  "Supervisor: teaching vs acquisition mode", "⛔",
     "No commits"),
    ("8",  "Early jump to doc generation", "🛠",
     "PR #150 — PRE-GENERATION INTAKE GATE in supervisor + market-intelligence prompts"),
    ("9",  "Re-run Q4 & Q5 after KB sync", "⛔",
     "Unblocked by #1; not yet re-run"),
    ("10", "Fri 4/17 follow-up meeting", "✅",
     "Happened; 20260417-kb-agent-prompt-comparison.md filed"),
    ("11", "Tone down RO-style risk framing", "🔄",
     "Stylistic — no discrete milestone"),
]

NEXT_QUEUE = [
    "#9 — Re-run Q4 + Q5 live. Fastest validation, unblocked, low cost.",
    "#2 + #8 — Open a fresh package where a PWS/IGCE is generated, confirm EAGLE asks for required intake facts in ONE batched question before generating, and does NOT ask the reconcile question when budget > IGCE.",
    "#5 — Inspect a generated IGCE from the deployed app, confirm §2.4 + §2.5 present.",
    "#3 — Ask EAGLE for a services acquisition on GSA Schedule, confirm T&M recommended (not LH) and FFP carve-outs offered.",
    "#1b — If #9 shows Q4/Q5 still wrong despite RFO KB, capture the agent trace and diagnose whether it's prompt drift, retrieval ranking, or a cited-citation-verification gap.",
]


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _header_row(table, labels: list[str]) -> None:
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, labels):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        _shade_cell(cell, "E7EEF7")


def build() -> None:
    doc = Document()

    # Title
    h = doc.add_heading("4/16 Output Review — Action Items Cheat Sheet", level=0)
    h.alignment = 0
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Updated 2026-04-23 · Source: 20260416-180100-meeting-eagle-output-review-v4.docx")
    sub_run.italic = True
    sub_run.font.size = Pt(9)

    # Legend
    doc.add_heading("Status legend", level=1)
    legend = doc.add_table(rows=1 + len(STATUS_LEGEND), cols=2)
    legend.style = "Light Grid Accent 1"
    legend.autofit = True
    _header_row(legend, ["Symbol", "Meaning"])
    for i, (sym, meaning) in enumerate(STATUS_LEGEND, start=1):
        legend.rows[i].cells[0].text = sym
        legend.rows[i].cells[1].text = meaning

    # Items
    doc.add_heading("Items", level=1)
    items = doc.add_table(rows=1 + len(ITEMS), cols=4)
    items.style = "Light Grid Accent 1"
    items.autofit = True
    _header_row(items, ["#", "Item", "Status", "Evidence"])

    status_fills = {
        "✅": "E8F3EB",  # validated — light green
        "🛠": "FFF4E1",  # patched — light amber
        "⛔": "FBE8E8",  # outstanding — light red
        "🔄": "EEE8F5",  # ongoing — light purple
    }

    for i, (num, item, status, evidence) in enumerate(ITEMS, start=1):
        row = items.rows[i].cells
        row[0].text = num
        row[1].text = item
        row[2].text = status
        row[3].text = evidence
        fill = status_fills.get(status)
        if fill:
            _shade_cell(row[2], fill)
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # Roll-up
    doc.add_heading("Roll-up", level=1)
    rollup = doc.add_paragraph()
    r = rollup.add_run("2 validated · 4 patched · 4 outstanding · 1 ongoing · 1 meeting-done")
    r.bold = True

    # Next queue
    doc.add_heading("Next validation queue", level=1)
    doc.add_paragraph(
        "Things to verify now that deploys are unblocked (PRs #150, #151, #152):"
    )
    for entry in NEXT_QUEUE:
        doc.add_paragraph(entry, style="List Number")

    doc.save(DEST)
    print(f"Wrote {DEST}")


if __name__ == "__main__":
    build()
