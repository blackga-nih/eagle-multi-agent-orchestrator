"""Tests for template_registry and template_service modules."""
import io
import pytest
from unittest.mock import MagicMock, patch

from app.template_registry import (
    TEMPLATE_REGISTRY,
    get_alternate_s3_keys,
    get_placeholder_map,
    get_template_mapping,
    get_template_s3_key,
    has_template,
    is_markdown_only,
    list_registered_doc_types,
)
from app.template_service import (
    DOCXPopulator,
    TemplateResult,
    TemplateService,
    XLSXPopulator,
    _cache_get,
    _cache_set,
    _MISS,
)
from app.xlsx_workbook_handlers import (
    detect_xlsx_handler_for_template_id,
    detect_xlsx_handler_for_workbook,
)


# ══════════════════════════════════════════════════════════════════════
#  Template Registry Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateRegistry:
    """Tests for template_registry module."""

    def test_has_template_for_registered_types(self):
        """Registered doc_types should return True."""
        assert has_template("sow") is True
        assert has_template("igce") is True
        assert has_template("market_research") is True
        assert has_template("justification") is True
        assert has_template("acquisition_plan") is True
        assert has_template("cor_certification") is True

    def test_has_template_for_markdown_only(self):
        """Markdown-only types should not have templates."""
        assert has_template("eval_criteria") is False
        assert has_template("security_checklist") is False
        assert has_template("section_508") is False
        assert has_template("contract_type_justification") is False

    def test_is_markdown_only(self):
        """Markdown-only types should be identified correctly."""
        assert is_markdown_only("eval_criteria") is True
        assert is_markdown_only("security_checklist") is True
        assert is_markdown_only("sow") is False
        assert is_markdown_only("igce") is False

    def test_get_template_mapping_returns_mapping(self):
        """Should return TemplateMapping for registered types."""
        mapping = get_template_mapping("sow")
        assert mapping is not None
        assert mapping.doc_type == "sow"
        assert mapping.file_type == "docx"
        assert mapping.s3_filename.endswith(".docx")

    def test_get_template_mapping_returns_none_for_unknown(self):
        """Should return None for unregistered types."""
        assert get_template_mapping("unknown_type") is None
        assert get_template_mapping("eval_criteria") is None

    def test_get_template_s3_key(self):
        """Should return full S3 key for registered types."""
        key = get_template_s3_key("sow")
        assert key is not None
        assert "approved/supervisor-core/essential-templates" in key
        assert key.endswith(".docx")

    def test_get_template_s3_key_none_for_unknown(self):
        """Should return None for unregistered types."""
        assert get_template_s3_key("unknown_type") is None

    def test_get_alternate_s3_keys(self):
        """Should return alternate keys for types with alternates."""
        alts = get_alternate_s3_keys("igce")
        assert isinstance(alts, list)
        assert len(alts) >= 1  # IGCE has Educational and Nonprofit alternates

    def test_get_placeholder_map(self):
        """Should return placeholder mapping for registered types."""
        placeholders = get_placeholder_map("sow")
        assert "title" in placeholders
        assert placeholders["title"] == "{{PROJECT_TITLE}}"

    def test_list_registered_doc_types(self):
        """Should list all registered doc types."""
        types = list_registered_doc_types()
        assert "sow" in types
        assert "igce" in types
        assert len(types) == len(TEMPLATE_REGISTRY)


# ══════════════════════════════════════════════════════════════════════
#  Template Service Cache Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateCache:
    """Tests for template caching."""

    def test_cache_miss_returns_sentinel(self):
        """Uncached items should return _MISS sentinel."""
        result = _cache_get("test-bucket", "nonexistent-key")
        assert result is _MISS

    def test_cache_set_and_get(self):
        """Cached items should be retrievable."""
        _cache_set("test-bucket", "test-key", b"test-data")
        result = _cache_get("test-bucket", "test-key")
        assert result == b"test-data"

    def test_cache_can_store_none(self):
        """Cache should distinguish between None (cached miss) and _MISS."""
        _cache_set("test-bucket", "none-key", None)
        result = _cache_get("test-bucket", "none-key")
        assert result is None  # Not _MISS


# ══════════════════════════════════════════════════════════════════════
#  DOCX Populator Tests
# ══════════════════════════════════════════════════════════════════════


class TestDOCXPopulator:
    """Tests for DOCXPopulator."""

    @pytest.fixture
    def sample_docx(self):
        """Create a minimal DOCX for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Project: {{PROJECT_TITLE}}")
        doc.add_paragraph("Description: {{DESCRIPTION}}")
        doc.add_paragraph("Tasks: {{TASKS}}")

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def test_populate_replaces_placeholders(self, sample_docx):
        """Placeholders should be replaced with data values."""
        data = {
            "title": "Cloud Migration",
            "description": "Migrate to AWS",
        }
        placeholder_map = {
            "title": "{{PROJECT_TITLE}}",
            "description": "{{DESCRIPTION}}",
        }

        result = DOCXPopulator.populate(sample_docx, data, placeholder_map)
        preview = DOCXPopulator.extract_text(result)

        assert "Cloud Migration" in preview
        assert "Migrate to AWS" in preview
        assert "{{PROJECT_TITLE}}" not in preview

    def test_populate_handles_lists(self, sample_docx):
        """List values should be converted to bullet points."""
        data = {
            "tasks": ["Task 1", "Task 2", "Task 3"],
        }
        placeholder_map = {
            "tasks": "{{TASKS}}",
        }

        result = DOCXPopulator.populate(sample_docx, data, placeholder_map)
        preview = DOCXPopulator.extract_text(result)

        assert "Task 1" in preview or "- Task 1" in preview

    def test_extract_text_returns_markdown(self, sample_docx):
        """Text extraction should return readable content."""
        preview = DOCXPopulator.extract_text(sample_docx)
        assert "Project:" in preview
        assert "Description:" in preview


# ══════════════════════════════════════════════════════════════════════
#  XLSX Populator Tests
# ══════════════════════════════════════════════════════════════════════


class TestXLSXPopulator:
    """Tests for XLSXPopulator."""

    @pytest.fixture
    def sample_xlsx(self):
        """Create a minimal XLSX for testing."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.title = "IGCE"
        ws["A1"] = "Project: {{PROJECT_TITLE}}"
        ws["A2"] = "Total: {{TOTAL_ESTIMATE}}"
        ws["A3"] = "{{LINE_ITEMS}}"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    @pytest.fixture
    def formula_xlsx(self):
        """Create an XLSX with a preserved formula for regression testing."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.title = "IGCE"
        ws["A1"] = 2
        ws["B1"] = 5
        ws["C1"] = "=A1*B1"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    @pytest.fixture
    def commercial_igce_xlsx(self):
        """Create a minimal workbook that matches the commercial IGCE layout."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        summary = wb.active
        summary.title = "IGCE"
        summary["G7"] = "=C7*E7"
        summary["G8"] = "=C8*E8"
        summary["G26"] = "=SUM(G7:G23)"
        summary["G38"] = "=SUM(E30:E37)"
        summary["H39"] = "=G26+G28+G38"

        for row in (11, 12, 13, 16, 17, 18, 21, 22, 23):
            summary[f"G{row}"] = f"=C{row}*E{row}"

        services = wb.create_sheet("IT Services")
        services["A5"] = "Expected Contract Type:"
        services["A6"] = "Period of Performance:"
        services["B6"] = "From:"
        services["D6"] = "To:"
        services["A12"] = "Senior Manager"
        services["D12"] = "=B12*C12"
        services["G12"] = "=E12*F12"
        services["J12"] = "=H12*I12"
        services["M12"] = "=K12*L12"
        services["P12"] = "=N12*O12"
        for row in (13, 14, 15, 16, 17, 18):
            services[f"D{row}"] = f"=B{row}*C{row}"
            services[f"G{row}"] = f"=E{row}*F{row}"
            services[f"J{row}"] = f"=H{row}*I{row}"
            services[f"M{row}"] = f"=K{row}*L{row}"
            services[f"P{row}"] = f"=N{row}*O{row}"

        goods = wb.create_sheet("IT Goods")
        goods["A5"] = "Expected Contract Type:"
        goods["A6"] = "Delivery Date"
        goods["G10"] = "=F10*E10"
        for row in (11, 12, 13, 14, 15, 16, 17):
            goods[f"G{row}"] = f"=F{row}*E{row}"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    @pytest.fixture
    def ige_products_xlsx(self):
        """Create a minimal workbook that matches the IGE for Products layout."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        sheet = wb.active
        sheet.title = "Sheet1"
        sheet["A1"] = "INDEPENDENT GOVERNMENT ESTIMATE (IGE) FOR PRODUCTS"
        sheet["B3"] = "Products (including software, licenses)"
        sheet["A4"] = "Item #"
        sheet["B4"] = "Manufacturer/Description/Part/Model Number*"
        sheet["C4"] = "Qty"
        sheet["D4"] = "Price"
        sheet["E4"] = "Extended Amount"
        for row in range(5, 10):
            sheet[f"E{row}"] = f"=SUM(C{row}*D{row})"
        sheet["B10"] = "Shipping (include as needed)"
        sheet["E10"] = "=SUM(C10*D10)"
        sheet["B11"] = "Installation (include as needed)"
        sheet["E11"] = "=SUM(C11*D11)"
        sheet["B12"] = "Training (include as needed)"
        sheet["E13"] = "=SUM(C13*D13)"
        sheet["B14"] = "  SUBTOTAL (Products)"
        sheet["E14"] = "=SUM(E5:E13)"
        sheet["B16"] = "TOTAL IGE"
        sheet["E16"] = "=E14"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    @pytest.fixture
    def ige_services_xlsx(self):
        """Create a minimal workbook matching the services catalog template layout."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        base = wb.active
        base.title = "Base Period"
        base["A1"] = "INDEPENDENT GOVERNMENT ESTIMATE (IGE) FOR SERVICES"
        base["B3"] = "Services"
        base["A4"] = "Brief Description of Service*"
        base["C4"] = "Period of Performance\n(Base)"
        base.merge_cells("C4:D4")
        for row in range(5, 14):
            base.merge_cells(f"A{row}:B{row}")
            base.merge_cells(f"C{row}:D{row}")
            base[f"E{row}"] = f"=SUM(C{row}*D{row})"
        base["B14"] = "  SUBTOTAL (Services)"
        base.merge_cells("C14:D14")
        base["B16"] = "TOTAL IGE"
        base["E16"] = "=E14"

        option = wb.create_sheet("Option Period One")
        option["A1"] = "INDEPENDENT GOVERNMENT COST ESTIMATE (IGCE)"
        option["B3"] = "Services"
        option["A4"] = "Brief Description of Service*"
        option["C4"] = "Period of Performance\n(Option Period 1)"
        option.merge_cells("C4:D4")
        for row in range(5, 14):
            option.merge_cells(f"A{row}:B{row}")
            option.merge_cells(f"C{row}:D{row}")
            option[f"E{row}"] = f"=SUM(C{row}*D{row})"
        option["B14"] = "  SUBTOTAL (Services)"
        option.merge_cells("C14:D14")
        option["B16"] = "TOTAL IGCE"
        option["E16"] = "=E14"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    def test_populate_replaces_placeholders(self, sample_xlsx):
        """Placeholders should be replaced with data values."""
        data = {
            "title": "Server Procurement",
            "total_estimate": "$150,000",
        }
        placeholder_map = {
            "title": "{{PROJECT_TITLE}}",
            "total_estimate": "{{TOTAL_ESTIMATE}}",
        }

        result = XLSXPopulator.populate(sample_xlsx, data, placeholder_map)
        preview = XLSXPopulator.extract_text(result)

        assert "Server Procurement" in preview
        assert "$150,000" in preview

    def test_populate_handles_line_items(self, sample_xlsx):
        """Line items array should be inserted into spreadsheet."""
        data = {
            "line_items": [
                {"description": "Server", "quantity": 10, "unit_price": 5000, "total": 50000},
                {"description": "Storage", "quantity": 5, "unit_price": 10000, "total": 50000},
            ],
        }
        placeholder_map = {
            "line_items": "{{LINE_ITEMS}}",
        }

        result = XLSXPopulator.populate(sample_xlsx, data, placeholder_map)

        # Verify by loading result
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result))
        ws = wb.active

        # Line items should be inserted starting at row 3
        assert ws["B3"].value == "Server"
        assert ws["B4"].value == "Storage"

    def test_extract_text_returns_markdown_table(self, sample_xlsx):
        """Text extraction should return markdown-formatted content."""
        preview = XLSXPopulator.extract_text(sample_xlsx)
        assert "IGCE" in preview  # Sheet title
        assert "|" in preview  # Table format

    def test_populate_preserves_existing_formulas(self, formula_xlsx):
        """XLSX population should not flatten workbook formulas."""
        result = XLSXPopulator.populate(formula_xlsx, {}, {})

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result), data_only=False)
        ws = wb.active

        assert ws["C1"].value == "=A1*B1"

    def test_populate_maps_commercial_igce_template(self, commercial_igce_xlsx):
        """Commercial IGCE workbook should use the template-aware mapper."""
        data = {
            "description": "Cloud Migration Support Services",
            "contract_type": "Firm-Fixed-Price",
            "period_of_performance": "12 months",
            "delivery_date": "2026-09-30",
            "prepared_by": "EAGLE System",
            "prepared_date": "2026-03-17",
            "line_items": [
                {
                    "description": "Cloud Architect",
                    "quantity": 2080,
                    "unit": "HR",
                    "unit_price": 175,
                    "total": 364000,
                },
                {
                    "description": "DevOps Engineer",
                    "quantity": 2080,
                    "unit": "HR",
                    "unit_price": 150,
                    "total": 312000,
                },
                {
                    "description": "AWS Licensing",
                    "quantity": 12,
                    "unit": "MO",
                    "unit_price": 15000,
                    "total": 180000,
                },
            ],
        }

        result = XLSXPopulator.populate(commercial_igce_xlsx, data, {})

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result), data_only=False)
        summary = wb["IGCE"]
        services = wb["IT Services"]
        goods = wb["IT Goods"]

        assert summary["A7"].value == "Cloud Architect"
        assert summary["C7"].value == 2080
        assert summary["E7"].value == 175
        assert summary["A8"].value == "DevOps Engineer"
        assert summary["E30"].value == 180000
        assert summary["G7"].value == "=C7*E7"

        assert services["B5"].value == "Firm-Fixed-Price"
        assert services["C6"].value == "12 months"
        assert services["A12"].value == "Cloud Architect"
        assert services["B12"].value == 2080
        assert services["C12"].value == 175
        assert services["E12"].value == 0
        assert services["D12"].value == "=B12*C12"

        assert goods["B5"].value == "Firm-Fixed-Price"
        assert goods["B6"].value == "2026-09-30"
        assert goods["A10"].value == "AWS Licensing"
        assert goods["E10"].value == 12
        assert goods["F10"].value == 15000
        assert goods["G10"].value == "=F10*E10"

    def test_populate_maps_total_estimate_when_line_items_missing(
        self,
        commercial_igce_xlsx,
    ):
        """Total estimate should still populate a usable workbook when items are absent."""
        result = XLSXPopulator.populate(
            commercial_igce_xlsx,
            {
                "description": "Seed Estimate",
                "total_estimate": "$12,500",
            },
            {},
        )

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result), data_only=False)
        summary = wb["IGCE"]
        goods = wb["IT Goods"]

        assert summary["A30"].value == "Seed Estimate"
        assert summary["E30"].value == 12500
        assert goods["A10"].value == "Seed Estimate"
        assert goods["E10"].value == 1
        assert goods["F10"].value == 12500

    def test_populate_maps_ige_products_template(self, ige_products_xlsx):
        result = XLSXPopulator.populate(
            ige_products_xlsx,
            {
                "line_items": [
                    {
                        "description": "Microscope",
                        "manufacturer": "Acme",
                        "part_number": "M-100",
                        "quantity": 2,
                        "unit_price": 5000,
                        "total": 10000,
                    },
                    {
                        "description": "Shipping",
                        "quantity": 1,
                        "unit_price": 400,
                        "total": 400,
                    },
                ]
            },
            {},
        )

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result), data_only=False)
        sheet = wb["Sheet1"]

        assert sheet["A5"].value == 1
        assert sheet["B5"].value == "Microscope / Acme / M-100"
        assert sheet["C5"].value == 2
        assert sheet["D5"].value == 5000
        assert sheet["E5"].value == "=SUM(C5*D5)"

        assert sheet["B10"].value == "Shipping (include as needed)"
        assert sheet["C10"].value == 1
        assert sheet["D10"].value == 400
        assert sheet["E10"].value == "=SUM(C10*D10)"

    def test_detects_registered_xlsx_handlers(self, commercial_igce_xlsx, ige_products_xlsx):
        from openpyxl import load_workbook

        commercial = load_workbook(io.BytesIO(commercial_igce_xlsx), data_only=False)
        products = load_workbook(io.BytesIO(ige_products_xlsx), data_only=False)

        assert detect_xlsx_handler_for_workbook(commercial).handler_id == "commercial_igce"
        assert detect_xlsx_handler_for_workbook(products).handler_id == "ige_products"
        assert (
            detect_xlsx_handler_for_template_id(
                "eagle-knowledge-base/approved/supervisor-core/essential-templates/4.a. IGE for Products.xlsx"
            ).handler_id
            == "ige_products"
        )

    def test_populate_maps_ige_services_catalog_template(self, ige_services_xlsx):
        result = XLSXPopulator.populate(
            ige_services_xlsx,
            {
                "line_items": [
                    {
                        "description": "Cloud Hosting",
                        "quantity": 12,
                        "unit_price": 5000,
                    }
                ],
                "option_period_one_items": [
                    {
                        "description": "Cloud Hosting Option 1",
                        "quantity": 12,
                        "unit_price": 5500,
                    }
                ],
            },
            {},
        )

        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result), data_only=False)
        base = wb["Base Period"]
        option = wb["Option Period One"]

        assert not any(str(rng) == "C5:D5" for rng in base.merged_cells.ranges)
        assert base["A5"].value == "Cloud Hosting"
        assert base["C5"].value == 12
        assert base["D5"].value == 5000
        assert base["E5"].value == "=SUM(C5*D5)"
        assert base["E14"].value == "=SUM(E5:E13)"

        assert option["A5"].value == "Cloud Hosting Option 1"
        assert option["C5"].value == 12
        assert option["D5"].value == 5500
        assert option["E5"].value == "=SUM(C5*D5)"


# ══════════════════════════════════════════════════════════════════════
#  Template Service Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateService:
    """Tests for TemplateService."""

    @pytest.fixture
    def mock_generators(self):
        """Create mock markdown generators."""
        return {
            "sow": lambda title, data: f"# SOW: {title}\n\nGenerated content",
            "eval_criteria": lambda title, data: f"# Eval Criteria: {title}\n\nMarkdown only",
        }

    def test_markdown_only_uses_generator(self, mock_generators):
        """Markdown-only doc types should use generator directly."""
        service = TemplateService("tenant1", "user1", mock_generators)
        result = service.generate_document("eval_criteria", "Test Eval", {})

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert result.file_type == "md"
        assert "Eval Criteria: Test Eval" in result.preview

    @patch("app.template_service.get_s3")
    def test_template_not_found_falls_back_to_markdown(self, mock_s3, mock_generators):
        """Missing S3 template should fall back to markdown."""
        from botocore.exceptions import ClientError

        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("tenant1", "user1", mock_generators)
        result = service.generate_document("sow", "Test SOW", {})

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert "SOW: Test SOW" in result.preview

    def test_no_generator_returns_error(self):
        """Missing generator should return error result."""
        service = TemplateService("tenant1", "user1", {})
        result = service._generate_markdown_fallback("unknown_type", "Test", {})

        assert result.success is False
        assert "No markdown generator" in result.error

    def test_template_result_dataclass(self):
        """TemplateResult should hold all expected fields."""
        result = TemplateResult(
            success=True,
            content=b"test content",
            preview="test preview",
            file_type="docx",
            source="s3_template",
            template_path="path/to/template.docx",
        )

        assert result.success is True
        assert result.content == b"test content"
        assert result.preview == "test preview"
        assert result.file_type == "docx"
        assert result.source == "s3_template"
        assert result.template_path == "path/to/template.docx"
        assert result.error is None


# ══════════════════════════════════════════════════════════════════════
#  Integration Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateServiceIntegration:
    """Integration tests for template service with real DOCX/XLSX."""

    @pytest.fixture
    def real_generators(self):
        """Import actual generators from create_document_support."""
        try:
            from app.tools.create_document_support import (
                _generate_acquisition_plan,
                _generate_cor_certification,
                _generate_contract_type_justification,
                _generate_eval_criteria,
                _generate_igce,
                _generate_justification,
                _generate_market_research,
                _generate_section_508,
                _generate_security_checklist,
                _generate_sow,
            )

            return {
                "sow": _generate_sow,
                "igce": _generate_igce,
                "market_research": _generate_market_research,
                "justification": _generate_justification,
                "acquisition_plan": _generate_acquisition_plan,
                "eval_criteria": _generate_eval_criteria,
                "security_checklist": _generate_security_checklist,
                "section_508": _generate_section_508,
                "cor_certification": _generate_cor_certification,
                "contract_type_justification": _generate_contract_type_justification,
            }
        except ImportError:
            pytest.skip("agentic_service generators not available")

    @pytest.mark.skip(reason="LLM output casing varies — not deterministic")
    @patch("app.template_service.get_s3")
    def test_full_fallback_flow(self, mock_s3, real_generators):
        """Test complete fallback flow with real generators."""
        from botocore.exceptions import ClientError

        # Simulate S3 not having the template
        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("test-tenant", "test-user", real_generators)

        # Test SOW generation with fallback
        result = service.generate_document(
            "sow",
            "Cloud Migration Services",
            {"description": "AWS migration", "period_of_performance": "12 months"},
        )

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert "Cloud Migration Services" in result.preview
        assert "STATEMENT OF WORK" in result.preview

    @patch("app.template_service.get_s3")
    def test_igce_fallback_with_line_items(self, mock_s3, real_generators):
        """Test IGCE generation with line items data."""
        from botocore.exceptions import ClientError

        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("test-tenant", "test-user", real_generators)

        result = service.generate_document(
            "igce",
            "Server Procurement IGCE",
            {
                "description": "Servers for data center",
                "total_estimate": "$500,000",
                "line_items": [
                    {"description": "Dell Servers", "quantity": 10, "unit_cost": "$25,000"},
                    {"description": "Storage Arrays", "quantity": 2, "unit_cost": "$100,000"},
                ],
            },
        )

        assert result.success is True
        assert "IGCE" in result.preview or "Cost Estimate" in result.preview
