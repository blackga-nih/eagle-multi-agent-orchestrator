"""Tests for enhanced upload pipeline — markdown persistence, auto-tagging, TTL extension.

Covers:
- Markdown sibling upload during document upload
- Extended TTL (24 hours)
- Auto-tagging during assign-to-package
- Doc type normalization in assign flow
"""

import io
import time
from decimal import Decimal
from unittest.mock import MagicMock, patch, ANY, call

import pytest
from fastapi.testclient import TestClient

from app.cognito_auth import UserContext


def _make_user(tenant_id: str = "test-tenant", user_id: str = "test-user") -> UserContext:
    return UserContext(
        user_id=user_id,
        tenant_id=tenant_id,
        email=f"{user_id}@example.com",
    )


@pytest.fixture()
def mock_user() -> UserContext:
    return _make_user()


@pytest.fixture()
def client(mock_user: UserContext):
    import app.main as _main
    _main.app.dependency_overrides[_main.get_user_from_header] = lambda: mock_user
    with TestClient(_main.app, raise_server_exceptions=False) as c:
        yield c
    _main.app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_MOCK_CLASSIFICATION = MagicMock()
_MOCK_CLASSIFICATION.doc_type = "sow"
_MOCK_CLASSIFICATION.to_dict.return_value = {
    "doc_type": "sow",
    "confidence": 0.95,
    "method": "filename",
    "suggested_title": "Statement of Work",
}


def _make_docx_bytes() -> bytes:
    """Create a minimal DOCX in memory."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Statement of Work", style="Heading 1")
    doc.add_paragraph("This is the scope of work for the project.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# TestUploadMarkdownPersistence
# ---------------------------------------------------------------------------


class TestUploadMarkdownPersistence:
    """Verify upload stores markdown sibling in S3 and metadata."""

    def test_upload_stores_markdown_s3_key(self, client):
        """Upload should store a .content.md sibling when conversion succeeds."""
        docx_bytes = _make_docx_bytes()

        mock_doc = {"document_id": "doc-md-1", "doc_type": "sow", "markdown_s3_key": "eagle/test/doc/v1/SOW-Test.docx.content.md"}

        with (
            patch("app.routers.documents.get_s3", return_value=MagicMock()),
            patch("app.routers.documents.classify_document", return_value=_MOCK_CLASSIFICATION),
            patch("app.routers.documents.extract_text_preview", return_value="preview"),
            patch("app.user_document_store.create_document", return_value=mock_doc) as mock_create,
        ):
            response = client.post(
                "/api/documents/upload",
                files={"file": ("SOW-Test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        assert response.status_code == 200
        # Verify create_document was called with a markdown_s3_key
        assert mock_create.called
        call_kwargs = mock_create.call_args.kwargs
        md_key = call_kwargs.get("markdown_s3_key")
        assert md_key is not None
        assert md_key.endswith(".content.md")

    def test_upload_plain_text_also_gets_markdown(self, client):
        """Plain text uploads should also produce a markdown sibling (passthrough)."""
        mock_doc = {"document_id": "doc-md-2", "doc_type": "sow"}

        with (
            patch("app.routers.documents.get_s3", return_value=MagicMock()),
            patch("app.routers.documents.classify_document", return_value=_MOCK_CLASSIFICATION),
            patch("app.routers.documents.extract_text_preview", return_value="plain text"),
            patch("app.user_document_store.create_document", return_value=mock_doc) as mock_create,
        ):
            response = client.post(
                "/api/documents/upload",
                files={"file": ("notes.txt", b"This is a test document.", "text/plain")},
            )

        assert response.status_code == 200
        call_kwargs = mock_create.call_args.kwargs
        assert call_kwargs.get("markdown_s3_key") is not None


# ---------------------------------------------------------------------------
# TestUploadExtendedTTL
# ---------------------------------------------------------------------------


class TestUploadExtendedTTL:
    """Verify upload metadata uses 24-hour TTL."""

    def test_ttl_is_24_hours(self):
        """_put_upload should set TTL to ~86400 seconds from now."""
        from app.routers.documents import _put_upload

        mock_table = MagicMock()
        with patch("app.routers.documents.get_table", return_value=mock_table):
            _put_upload("test-tenant", "upload-123", {"key": "value"})

        assert mock_table.put_item.called
        item = mock_table.put_item.call_args[1]["Item"]
        ttl = item["ttl"]
        expected_min = int(time.time()) + 86400 - 60  # within 60s tolerance
        expected_max = int(time.time()) + 86400 + 60
        assert expected_min <= ttl <= expected_max, f"TTL {ttl} not within 24h range"


# ---------------------------------------------------------------------------
# TestAssignDocTypeNormalization
# ---------------------------------------------------------------------------


class TestAssignDocTypeNormalization:
    """Verify assign-to-package normalizes doc_type."""

    def test_hyphenated_doc_type_normalized(self, client):
        """Request with 'acquisition-plan' should be normalized to 'acquisition_plan'."""
        upload_meta = {
            "tenant_id": "test-tenant",
            "user_id": "test-user",
            "s3_bucket": "test-bucket",
            "s3_key": "eagle/test/uploads/abc/doc.docx",
            "filename": "AP-Test.docx",
            "original_filename": "AP-Test.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "classification": {"doc_type": "acquisition_plan", "confidence": 0.95, "method": "filename", "suggested_title": "AP Test"},
        }

        mock_result = MagicMock()
        mock_result.success = True
        mock_result.to_dict.return_value = {
            "success": True,
            "doc_type": "acquisition_plan",
            "version": 1,
        }

        with (
            patch("app.routers.documents._get_upload", return_value=upload_meta),
            patch("app.routers.documents.get_package", return_value={"package_id": "pkg-1"}),
            patch("boto3.client") as mock_boto3,
            patch("app.routers.documents.create_package_document_version", return_value=mock_result) as mock_create,
            patch("app.routers.documents._delete_upload"),
            patch("app.tag_computation.compute_document_tags", return_value=["phase:planning"]),
            patch("app.tag_computation.compute_far_tags_from_template", return_value=["FAR 7.105"]),
            patch("app.tag_computation.compute_completeness_pct", return_value=0.0),
        ):
            mock_s3 = MagicMock()
            mock_s3.get_object.return_value = {"Body": MagicMock(read=lambda: b"content")}
            mock_boto3.return_value = mock_s3

            response = client.post(
                "/api/documents/upload-123/assign-to-package",
                json={"package_id": "pkg-1", "doc_type": "acquisition-plan"},
            )

        assert response.status_code == 200
        # Verify create_package_document_version was called with normalized type
        assert mock_create.called
        call_kwargs = mock_create.call_args[1]
        assert call_kwargs["doc_type"] == "acquisition_plan"


# ---------------------------------------------------------------------------
# TestAssignAutoTags
# ---------------------------------------------------------------------------


class TestAssignAutoTags:
    """Verify assign-to-package computes and passes tags."""

    def test_assign_passes_tags_to_create(self, client):
        """Tags computed from tag_computation should be passed to create_package_document_version."""
        upload_meta = {
            "tenant_id": "test-tenant",
            "user_id": "test-user",
            "s3_bucket": "test-bucket",
            "s3_key": "eagle/test/uploads/abc/sow.docx",
            "filename": "sow.docx",
            "original_filename": "SOW-Test.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "classification": {"doc_type": "sow", "confidence": 0.95, "method": "filename", "suggested_title": "SOW Test"},
            "markdown_s3_key": "eagle/test/uploads/abc/sow.docx.parsed.md",
        }

        mock_result = MagicMock()
        mock_result.success = True
        mock_result.to_dict.return_value = {"success": True, "doc_type": "sow", "version": 1}

        expected_sys_tags = ["phase:planning", "doc_type:sow"]
        expected_far_tags = ["FAR 37.6"]

        with (
            patch("app.routers.documents._get_upload", return_value=upload_meta),
            patch("app.routers.documents.get_package", return_value={"package_id": "pkg-1"}),
            patch("boto3.client") as mock_boto3,
            patch("app.routers.documents.create_package_document_version", return_value=mock_result) as mock_create,
            patch("app.routers.documents._delete_upload"),
            patch("app.tag_computation.compute_document_tags", return_value=expected_sys_tags),
            patch("app.tag_computation.compute_far_tags_from_template", return_value=expected_far_tags),
            patch("app.tag_computation.compute_completeness_pct", return_value=42.5),
        ):
            mock_s3 = MagicMock()
            mock_s3.get_object.return_value = {"Body": MagicMock(read=lambda: b"# SOW content")}
            mock_boto3.return_value = mock_s3

            response = client.post(
                "/api/documents/upload-456/assign-to-package",
                json={"package_id": "pkg-1"},
            )

        assert response.status_code == 200
        assert mock_create.called
        call_kwargs = mock_create.call_args[1]
        assert call_kwargs["system_tags"] == expected_sys_tags
        assert call_kwargs["far_tags"] == expected_far_tags
        assert call_kwargs["completeness_pct"] == 42.5
        assert call_kwargs["original_filename"] == "SOW-Test.docx"
        assert call_kwargs["markdown_content"] is not None
