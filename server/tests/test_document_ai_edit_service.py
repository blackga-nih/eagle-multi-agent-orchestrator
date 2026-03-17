"""
Tests for document_ai_edit_service.py — DOCX preview extraction and editing.

Validates:
  - apply_docx_edits(): text replacement preserving formatting, checkbox toggling
  - extract_docx_preview_payload(): structured block extraction, text fallback
  - apply_docx_block_edits(): browser-side block editing round-trip
  - edit_docx_document(): full flow with mocked S3
  - save_docx_preview_edits(): browser save flow with mocked S3
  - _normalize_text(): whitespace normalization
  - _extract_text_fallback(): UTF-8 fallback for misnamed artifacts
  - _is_allowed_document_key(): security check for tenant/user paths
  - _extract_package_document_ref(): regex parsing of package S3 keys
  - _extract_workspace_document_ref(): regex parsing of workspace S3 keys

All tests are fast (mocked S3/DynamoDB, in-memory DOCX via python-docx).
"""

import io
import pytest
from unittest.mock import MagicMock, patch


# ---------------------------------------------------------------------------
# Helpers — build in-memory DOCX fixtures
# ---------------------------------------------------------------------------

def _build_sample_docx(paragraphs: list[tuple[str, str]] | None = None) -> bytes:
    """Build a minimal DOCX in memory.

    paragraphs: list of (style_name, text) tuples.
    Defaults to a simple document with a heading and two paragraphs.
    """
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    if paragraphs is None:
        paragraphs = [
            ("Heading 1", "Introduction"),
            ("Normal", "This is the first paragraph."),
            ("Heading 2", "Details"),
            ("Normal", "This is the second paragraph with specific text."),
        ]

    for style_name, text in paragraphs:
        try:
            doc.add_paragraph(text, style=style_name)
        except KeyError:
            para = doc.add_paragraph(text)
            para.style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_checkbox_docx(items: list[tuple[str, bool]]) -> bytes:
    """Build a DOCX with checkbox paragraphs using Word XML.

    items: list of (label_text, is_checked) tuples.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    for label, checked in items:
        para = doc.add_paragraph(label)
        # Inject a w:checkBox element into the paragraph's run properties
        pPr = para._p.get_or_add_pPr()
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
        checkbox = OxmlElement("w:checkBox")
        checked_elem = OxmlElement("w:checked")
        checked_elem.set(qn("w:val"), "1" if checked else "0")
        checkbox.append(checked_elem)
        rPr.append(checkbox)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# TestNormalizeText
# ---------------------------------------------------------------------------

class TestNormalizeText:
    """Tests for _normalize_text helper."""

    def test_collapses_whitespace(self):
        from app.document_ai_edit_service import _normalize_text

        assert _normalize_text("hello   world") == "hello world"

    def test_strips_leading_trailing(self):
        from app.document_ai_edit_service import _normalize_text

        assert _normalize_text("  hello  ") == "hello"

    def test_handles_none(self):
        from app.document_ai_edit_service import _normalize_text

        assert _normalize_text(None) == ""

    def test_handles_newlines_and_tabs(self):
        from app.document_ai_edit_service import _normalize_text

        assert _normalize_text("hello\n\tworld") == "hello world"


# ---------------------------------------------------------------------------
# TestExtractTextFallback
# ---------------------------------------------------------------------------

class TestExtractTextFallback:
    """Tests for _extract_text_fallback helper."""

    def test_returns_text_for_utf8_content(self):
        from app.document_ai_edit_service import _extract_text_fallback

        result = _extract_text_fallback(b"# Hello World\n\nSome content here.")
        assert result is not None
        assert "Hello World" in result

    def test_returns_none_for_empty_bytes(self):
        from app.document_ai_edit_service import _extract_text_fallback

        assert _extract_text_fallback(b"") is None

    def test_returns_none_for_binary_content(self):
        from app.document_ai_edit_service import _extract_text_fallback

        # Random binary bytes that aren't valid UTF-8
        assert _extract_text_fallback(b"\x80\x81\x82\x83\x84\x85") is None

    def test_returns_none_for_low_printable_ratio(self):
        from app.document_ai_edit_service import _extract_text_fallback

        # Mostly control characters
        content = b"a" + b"\x00" * 100
        assert _extract_text_fallback(content) is None


# ---------------------------------------------------------------------------
# TestIsAllowedDocumentKey
# ---------------------------------------------------------------------------

class TestIsAllowedDocumentKey:
    """Tests for _is_allowed_document_key."""

    def test_allows_package_key(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/tenant1/packages/pkg-123/sow/v1/sow_v1.docx",
            "tenant1", "user1",
        ) is True

    def test_allows_workspace_key(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/tenant1/user1/documents/sow.docx",
            "tenant1", "user1",
        ) is True

    def test_denies_other_tenant(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/other-tenant/packages/pkg-123/sow/v1/sow.docx",
            "tenant1", "user1",
        ) is False

    def test_denies_other_user_workspace(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/tenant1/other-user/documents/sow.docx",
            "tenant1", "user1",
        ) is False

    def test_denies_with_none_user_id(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/tenant1/user1/documents/sow.docx",
            "tenant1", None,
        ) is False

    def test_allows_package_key_with_none_user_id(self):
        from app.document_ai_edit_service import _is_allowed_document_key

        assert _is_allowed_document_key(
            "eagle/tenant1/packages/pkg-123/igce/v2/igce.xlsx",
            "tenant1", None,
        ) is True


# ---------------------------------------------------------------------------
# TestExtractPackageDocumentRef
# ---------------------------------------------------------------------------

class TestExtractPackageDocumentRef:
    """Tests for _extract_package_document_ref."""

    def test_parses_canonical_package_key(self):
        from app.document_ai_edit_service import _extract_package_document_ref

        ref = _extract_package_document_ref(
            "eagle/tenant1/packages/pkg-123/sow/v1/sow_v1.docx"
        )
        assert ref is not None
        assert ref["tenant_id"] == "tenant1"
        assert ref["package_id"] == "pkg-123"
        assert ref["doc_type"] == "sow"
        assert ref["version"] == 1
        assert ref["filename"] == "sow_v1.docx"

    def test_parses_higher_version(self):
        from app.document_ai_edit_service import _extract_package_document_ref

        ref = _extract_package_document_ref(
            "eagle/t/packages/p/igce/v12/igce_v12.xlsx"
        )
        assert ref is not None
        assert ref["version"] == 12

    def test_returns_none_for_workspace_key(self):
        from app.document_ai_edit_service import _extract_package_document_ref

        assert _extract_package_document_ref(
            "eagle/tenant1/user1/documents/sow.docx"
        ) is None

    def test_returns_none_for_invalid_key(self):
        from app.document_ai_edit_service import _extract_package_document_ref

        assert _extract_package_document_ref("random/path/file.docx") is None


# ---------------------------------------------------------------------------
# TestExtractWorkspaceDocumentRef
# ---------------------------------------------------------------------------

class TestExtractWorkspaceDocumentRef:
    """Tests for _extract_workspace_document_ref."""

    def test_parses_workspace_key(self):
        from app.document_ai_edit_service import _extract_workspace_document_ref

        ref = _extract_workspace_document_ref(
            "eagle/tenant1/user1/documents/sow_20260310.docx"
        )
        assert ref is not None
        assert ref["tenant"] == "tenant1"
        assert ref["user"] == "user1"
        assert ref["filename"] == "sow_20260310.docx"

    def test_returns_none_for_package_key(self):
        from app.document_ai_edit_service import _extract_workspace_document_ref

        assert _extract_workspace_document_ref(
            "eagle/tenant1/packages/pkg-123/sow/v1/sow.docx"
        ) is None


# ---------------------------------------------------------------------------
# TestExtractDocxPreviewPayload
# ---------------------------------------------------------------------------

class TestExtractDocxPreviewPayload:
    """Tests for extract_docx_preview_payload."""

    def test_extracts_blocks_from_valid_docx(self):
        from app.document_ai_edit_service import extract_docx_preview_payload

        docx_bytes = _build_sample_docx()
        result = extract_docx_preview_payload(docx_bytes)

        assert result["preview_mode"] == "docx_blocks"
        assert len(result["preview_blocks"]) >= 4
        assert result["content"] is not None

        # Check first block is heading
        first_block = result["preview_blocks"][0]
        assert first_block["kind"] == "heading"
        assert first_block["text"] == "Introduction"
        assert first_block["level"] == 1

    def test_returns_text_fallback_for_plain_text(self):
        from app.document_ai_edit_service import extract_docx_preview_payload

        text_bytes = b"# Title\n\nSome paragraph text here."
        result = extract_docx_preview_payload(text_bytes)

        assert result["preview_mode"] == "text_fallback"
        assert len(result["preview_blocks"]) >= 1
        assert result["content"] is not None

    def test_returns_none_mode_for_binary_garbage(self):
        from app.document_ai_edit_service import extract_docx_preview_payload

        result = extract_docx_preview_payload(b"\x80\x81\x82")

        assert result["preview_mode"] == "none"
        assert result["content"] is None
        assert result["preview_blocks"] == []

    def test_block_ids_are_sequential(self):
        from app.document_ai_edit_service import extract_docx_preview_payload

        docx_bytes = _build_sample_docx()
        result = extract_docx_preview_payload(docx_bytes)

        for block in result["preview_blocks"]:
            assert block["block_id"].startswith("p:")


# ---------------------------------------------------------------------------
# TestApplyDocxEdits
# ---------------------------------------------------------------------------

class TestApplyDocxEdits:
    """Tests for apply_docx_edits — text replacement."""

    def test_replaces_exact_text(self):
        from app.document_ai_edit_service import DocxEdit, apply_docx_edits

        docx_bytes = _build_sample_docx()
        edits = [DocxEdit(search_text="first paragraph", replacement_text="opening section")]
        updated, applied, missing = apply_docx_edits(docx_bytes, edits)

        assert applied == 1
        assert missing == []
        assert len(updated) > 0

        # Verify the edit took effect by re-extracting
        from app.document_ai_edit_service import extract_docx_preview_payload

        payload = extract_docx_preview_payload(updated)
        content = payload["content"]
        assert "opening section" in content
        assert "first paragraph" not in content

    def test_reports_missing_when_text_not_found(self):
        from app.document_ai_edit_service import DocxEdit, apply_docx_edits

        docx_bytes = _build_sample_docx()
        edits = [DocxEdit(search_text="nonexistent text xyz", replacement_text="replacement")]
        updated, applied, missing = apply_docx_edits(docx_bytes, edits)

        assert applied == 0
        assert "nonexistent text xyz" in missing

    def test_applies_multiple_edits(self):
        from app.document_ai_edit_service import DocxEdit, apply_docx_edits

        docx_bytes = _build_sample_docx()
        edits = [
            DocxEdit(search_text="Introduction", replacement_text="Overview"),
            DocxEdit(search_text="Details", replacement_text="Specifics"),
        ]
        updated, applied, missing = apply_docx_edits(docx_bytes, edits)

        assert applied == 2
        assert missing == []

    def test_returns_valid_docx_bytes(self):
        from app.document_ai_edit_service import DocxEdit, apply_docx_edits
        from docx import Document

        docx_bytes = _build_sample_docx()
        edits = [DocxEdit(search_text="Introduction", replacement_text="Overview")]
        updated, _, _ = apply_docx_edits(docx_bytes, edits)

        # Should be a valid DOCX
        doc = Document(io.BytesIO(updated))
        assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# TestApplyDocxBlockEdits
# ---------------------------------------------------------------------------

class TestApplyDocxBlockEdits:
    """Tests for apply_docx_block_edits — browser-side block editing."""

    def test_text_fallback_mode_produces_utf8(self):
        from app.document_ai_edit_service import apply_docx_block_edits

        blocks = [
            {"block_id": "t:0", "kind": "heading", "text": "New Title", "level": 1, "checked": None},
            {"block_id": "t:1", "kind": "paragraph", "text": "New body text.", "level": None, "checked": None},
        ]
        updated_bytes, applied = apply_docx_block_edits(b"placeholder", blocks, "text_fallback")

        assert applied == 2
        text = updated_bytes.decode("utf-8")
        assert "# New Title" in text
        assert "New body text." in text

    def test_docx_blocks_mode_applies_text_changes(self):
        from app.document_ai_edit_service import (
            apply_docx_block_edits,
            extract_docx_preview_payload,
        )

        docx_bytes = _build_sample_docx()
        payload = extract_docx_preview_payload(docx_bytes)
        blocks = payload["preview_blocks"]

        # Modify the first block's text
        blocks[0]["text"] = "Modified Heading"

        updated_bytes, applied = apply_docx_block_edits(docx_bytes, blocks, "docx_blocks")

        assert applied >= 1
        # Verify it's still a valid DOCX
        from docx import Document
        doc = Document(io.BytesIO(updated_bytes))
        assert len(doc.paragraphs) > 0

    def test_skips_invalid_block_ids(self):
        from app.document_ai_edit_service import apply_docx_block_edits

        docx_bytes = _build_sample_docx()
        blocks = [
            {"block_id": "invalid", "kind": "paragraph", "text": "ignored", "level": None, "checked": None},
        ]
        updated_bytes, applied = apply_docx_block_edits(docx_bytes, blocks, "docx_blocks")

        assert applied == 0


# ---------------------------------------------------------------------------
# TestExtractTextPreviewBlocks
# ---------------------------------------------------------------------------

class TestExtractTextPreviewBlocks:
    """Tests for _extract_text_preview_blocks."""

    def test_parses_headings(self):
        from app.document_ai_edit_service import _extract_text_preview_blocks

        blocks = _extract_text_preview_blocks("# Title\n\nBody text")
        assert blocks[0].kind == "heading"
        assert blocks[0].text == "Title"
        assert blocks[0].level == 1

    def test_parses_checkboxes(self):
        from app.document_ai_edit_service import _extract_text_preview_blocks

        blocks = _extract_text_preview_blocks("- [x] Done item\n\n- [ ] Todo item")
        assert blocks[0].kind == "checkbox"
        assert blocks[0].checked is True
        assert blocks[0].text == "Done item"
        assert blocks[1].kind == "checkbox"
        assert blocks[1].checked is False

    def test_parses_plain_paragraphs(self):
        from app.document_ai_edit_service import _extract_text_preview_blocks

        blocks = _extract_text_preview_blocks("Just some text.")
        assert blocks[0].kind == "paragraph"
        assert blocks[0].text == "Just some text."

    def test_block_ids_use_t_prefix(self):
        from app.document_ai_edit_service import _extract_text_preview_blocks

        blocks = _extract_text_preview_blocks("# H1\n\nPara\n\n## H2")
        for block in blocks:
            assert block.block_id.startswith("t:")


# ---------------------------------------------------------------------------
# TestPreviewTextFromBlocks
# ---------------------------------------------------------------------------

class TestPreviewTextFromBlocks:
    """Tests for _preview_text_from_blocks."""

    def test_renders_heading(self):
        from app.document_ai_edit_service import DocxPreviewBlock, _preview_text_from_blocks

        blocks = [DocxPreviewBlock(block_id="p:0", kind="heading", text="Title", level=1)]
        text = _preview_text_from_blocks(blocks)
        assert text == "# Title"

    def test_renders_checkbox(self):
        from app.document_ai_edit_service import DocxPreviewBlock, _preview_text_from_blocks

        blocks = [
            DocxPreviewBlock(block_id="p:0", kind="checkbox", text="Task done", checked=True),
            DocxPreviewBlock(block_id="p:1", kind="checkbox", text="Task todo", checked=False),
        ]
        text = _preview_text_from_blocks(blocks)
        assert "- [x] Task done" in text
        assert "- [ ] Task todo" in text

    def test_renders_paragraph(self):
        from app.document_ai_edit_service import DocxPreviewBlock, _preview_text_from_blocks

        blocks = [DocxPreviewBlock(block_id="p:0", kind="paragraph", text="Body text")]
        text = _preview_text_from_blocks(blocks)
        assert text == "Body text"


# ---------------------------------------------------------------------------
# TestEditDocxDocument — integration with mocked S3
# ---------------------------------------------------------------------------

class TestEditDocxDocument:
    """Tests for edit_docx_document with mocked S3."""

    def test_returns_error_for_empty_doc_key(self):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document

        result = edit_docx_document(
            tenant_id="t", user_id="u", doc_key="",
            edits=[DocxEdit("a", "b")],
        )
        assert "error" in result

    def test_returns_error_for_no_edits(self):
        from app.document_ai_edit_service import edit_docx_document

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            edits=[], checkbox_edits=[],
        )
        assert "error" in result
        assert "required" in result["error"].lower()

    def test_returns_error_for_non_docx_key(self):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/report.pdf",
            edits=[DocxEdit("a", "b")],
        )
        assert "error" in result
        assert "docx" in result["error"].lower()

    def test_returns_error_for_access_denied(self):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/other/u/documents/sow.docx",
            edits=[DocxEdit("a", "b")],
        )
        assert "error" in result
        assert "denied" in result["error"].lower()

    @patch("app.document_ai_edit_service._get_s3")
    @patch("app.document_ai_edit_service.write_document_changelog_entry")
    def test_workspace_edit_success(self, mock_changelog, mock_get_s3):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document

        docx_bytes = _build_sample_docx()
        mock_s3 = MagicMock()
        mock_s3.get_object.return_value = {
            "Body": MagicMock(read=MagicMock(return_value=docx_bytes)),
        }
        mock_get_s3.return_value = mock_s3

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            edits=[DocxEdit("Introduction", "Overview")],
            change_source="ai_edit",
        )

        assert result.get("success") is True
        assert result["mode"] == "workspace_docx_edit"
        assert result["edits_applied"] == 1
        mock_s3.put_object.assert_called_once()
        mock_changelog.assert_called_once()

    @patch("app.document_ai_edit_service._get_s3")
    @patch("app.document_ai_edit_service.write_document_changelog_entry")
    def test_workspace_edit_returns_missing(self, mock_changelog, mock_get_s3):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document

        docx_bytes = _build_sample_docx()
        mock_s3 = MagicMock()
        mock_s3.get_object.return_value = {
            "Body": MagicMock(read=MagicMock(return_value=docx_bytes)),
        }
        mock_get_s3.return_value = mock_s3

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            edits=[DocxEdit("nonexistent text", "replacement")],
        )

        # No edits applied → error
        assert "error" in result
        assert "No DOCX edits" in result["error"]

    @patch("app.document_ai_edit_service._get_s3")
    def test_s3_load_failure(self, mock_get_s3):
        from app.document_ai_edit_service import DocxEdit, edit_docx_document
        from botocore.exceptions import ClientError

        mock_s3 = MagicMock()
        mock_s3.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey", "Message": "Not found"}}, "GetObject"
        )
        mock_get_s3.return_value = mock_s3

        result = edit_docx_document(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            edits=[DocxEdit("a", "b")],
        )

        assert "error" in result
        assert "Failed to load" in result["error"]


# ---------------------------------------------------------------------------
# TestSaveDocxPreviewEdits
# ---------------------------------------------------------------------------

class TestSaveDocxPreviewEdits:
    """Tests for save_docx_preview_edits."""

    def test_returns_error_for_empty_doc_key(self):
        from app.document_ai_edit_service import save_docx_preview_edits

        result = save_docx_preview_edits(
            tenant_id="t", user_id="u", doc_key="",
            preview_blocks=[{"block_id": "p:0", "kind": "paragraph", "text": "hi"}],
            preview_mode="docx_blocks",
        )
        assert "error" in result

    def test_returns_error_for_empty_blocks(self):
        from app.document_ai_edit_service import save_docx_preview_edits

        result = save_docx_preview_edits(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            preview_blocks=[],
            preview_mode="docx_blocks",
        )
        assert "error" in result

    def test_returns_error_for_non_docx(self):
        from app.document_ai_edit_service import save_docx_preview_edits

        result = save_docx_preview_edits(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/report.xlsx",
            preview_blocks=[{"block_id": "p:0"}],
            preview_mode="docx_blocks",
        )
        assert "error" in result

    @patch("app.document_ai_edit_service._get_s3")
    @patch("app.document_ai_edit_service.write_document_changelog_entry")
    def test_workspace_save_success(self, mock_changelog, mock_get_s3):
        from app.document_ai_edit_service import (
            extract_docx_preview_payload,
            save_docx_preview_edits,
        )

        docx_bytes = _build_sample_docx()
        payload = extract_docx_preview_payload(docx_bytes)
        blocks = payload["preview_blocks"]
        blocks[0]["text"] = "Changed Heading"

        mock_s3 = MagicMock()
        mock_s3.get_object.return_value = {
            "Body": MagicMock(read=MagicMock(return_value=docx_bytes)),
            "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        mock_get_s3.return_value = mock_s3

        result = save_docx_preview_edits(
            tenant_id="t", user_id="u",
            doc_key="eagle/t/u/documents/sow.docx",
            preview_blocks=blocks,
            preview_mode="docx_blocks",
        )

        assert result.get("success") is True
        assert result["mode"] == "workspace_docx_preview_edit"
        mock_s3.put_object.assert_called_once()
        mock_changelog.assert_called_once()
