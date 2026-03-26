"""Tests for binary-safe document endpoint behavior."""

import io
import os
import sys
from datetime import datetime
from types import ModuleType
from unittest.mock import MagicMock, patch

from fastapi import APIRouter
from fastapi.testclient import TestClient

# Ensure server/ is on sys.path so "app.main" resolves
_server_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), ".."))
if _server_dir not in sys.path:
    sys.path.insert(0, _server_dir)


ENV_PATCH = {
    "REQUIRE_AUTH": "false",
    "DEV_MODE": "true",
    "USE_BEDROCK": "false",
    "COGNITO_USER_POOL_ID": "us-east-1_test",
    "COGNITO_CLIENT_ID": "test-client",
    "EAGLE_SESSIONS_TABLE": "eagle",
    "USE_PERSISTENT_SESSIONS": "false",
    "S3_BUCKET": "test-bucket",
}


def _build_docx_bytes(text: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Statement of Work", level=1)
    doc.add_paragraph(text)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _build_xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "IGCE"
    ws["A1"] = "Item"
    ws["B1"] = "Estimate"
    ws["A2"] = "Olympus CK2"
    ws["B2"] = 13850
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _build_s3_mock() -> MagicMock:
    s3 = MagicMock()
    state = {
        "docx_bytes": _build_docx_bytes("Existing scope paragraph"),
        "xlsx_bytes": _build_xlsx_bytes(),
    }

    def get_object(*, Bucket, Key):
        # Sidecar markdown requests (.content.md) should 404 so binary
        # preview extraction runs instead.
        if Key.endswith(".content.md"):
            from botocore.exceptions import ClientError

            raise ClientError(
                {"Error": {"Code": "NoSuchKey", "Message": "Not found"}},
                "GetObject",
            )
        if Key.endswith(".docx"):
            return {
                "Body": io.BytesIO(state["docx_bytes"]),
                "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "ContentLength": 1024,
                "LastModified": datetime(2026, 3, 11, 12, 0, 0),
            }
        if Key.endswith(".xlsx"):
            return {
                "Body": io.BytesIO(state["xlsx_bytes"]),
                "ContentType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "ContentLength": 1024,
                "LastModified": datetime(2026, 3, 11, 12, 2, 0),
            }
        return {
            "Body": io.BytesIO(b"# Test markdown\n\nhello"),
            "ContentType": "text/markdown; charset=utf-8",
            "ContentLength": 24,
            "LastModified": datetime(2026, 3, 11, 12, 5, 0),
        }

    s3.get_object.side_effect = get_object
    s3.generate_presigned_url.return_value = "https://signed.example/document"

    def put_object(*, Bucket, Key, Body, ContentType):
        if Key.endswith(".docx"):
            state["docx_bytes"] = Body
        elif Key.endswith(".xlsx"):
            state["xlsx_bytes"] = Body
        return {"ResponseMetadata": {"HTTPStatusCode": 200}}

    s3.put_object.side_effect = put_object
    return s3


def _load_app():
    with patch.dict(os.environ, ENV_PATCH, clear=False):
        fake_strands = ModuleType("app.strands_agentic_service")
        fake_streaming_routes = ModuleType("app.streaming_routes")

        async def _mock_sdk_query(*args, **kwargs):
            if False:
                yield None

        async def _mock_sdk_query_streaming(*args, **kwargs):
            if False:
                yield None

        fake_strands.sdk_query = _mock_sdk_query
        fake_strands.sdk_query_streaming = _mock_sdk_query_streaming
        fake_strands.MODEL = "test-model"
        fake_strands.EAGLE_TOOLS = []
        fake_streaming_routes.create_streaming_router = lambda *args, **kwargs: APIRouter()

        with patch.dict(
            sys.modules,
            {
                "app.strands_agentic_service": fake_strands,
                "app.streaming_routes": fake_streaming_routes,
            },
        ):
            import importlib

            for _mod in ("app.main", "app.changelog_store", "app.cognito_auth", "app.document_ai_edit_service", "app.spreadsheet_edit_service"):
                sys.modules.pop(_mod, None)

            import app.main as main_module

            importlib.reload(main_module)

            # Force DEV_MODE=true on the reloaded cognito_auth module
            # (.env loads with override=True during reload, clobbering the patch)
            import app.cognito_auth as _auth
            _auth.DEV_MODE = True

            return main_module


def _auth_header() -> dict:
    """Generate a Bearer token header for test requests."""
    from app.cognito_auth import generate_test_token
    token = generate_test_token(user_id="dev-user", tenant_id="dev-tenant")
    return {"Authorization": f"Bearer {token}"}


def _load_app_without_docx_service():
    with patch.dict(os.environ, ENV_PATCH, clear=False):
        fake_strands = ModuleType("app.strands_agentic_service")
        fake_streaming_routes = ModuleType("app.streaming_routes")
        fake_docx_service = ModuleType("app.document_ai_edit_service")

        async def _mock_sdk_query(*args, **kwargs):
            if False:
                yield None

        async def _mock_sdk_query_streaming(*args, **kwargs):
            if False:
                yield None

        fake_strands.sdk_query = _mock_sdk_query
        fake_strands.sdk_query_streaming = _mock_sdk_query_streaming
        fake_strands.MODEL = "test-model"
        fake_strands.EAGLE_TOOLS = []
        fake_streaming_routes.create_streaming_router = lambda *args, **kwargs: APIRouter()
        fake_docx_service.extract_docx_preview_payload = lambda raw_bytes: {
            "content": None,
            "preview_blocks": [],
            "preview_mode": "none",
        }
        fake_docx_service.save_docx_preview_edits = lambda *args, **kwargs: {
            "error": "DOCX editing unavailable in test stub",
        }

        with patch.dict(
            sys.modules,
            {
                "app.strands_agentic_service": fake_strands,
                "app.streaming_routes": fake_streaming_routes,
                "app.document_ai_edit_service": fake_docx_service,
            },
        ):
            import importlib

            for _mod in ("app.main", "app.changelog_store", "app.cognito_auth", "app.document_ai_edit_service", "app.spreadsheet_edit_service"):
                sys.modules.pop(_mod, None)

            import app.main as main_module

            importlib.reload(main_module)

            import app.cognito_auth as _auth
            _auth.DEV_MODE = True

            return main_module


def test_get_document_returns_binary_metadata_for_canonical_package_doc():
    main_module = _load_app()
    app = main_module.app
    s3 = _build_s3_mock()

    with patch("boto3.client", return_value=s3), patch(
        "app.routers.documents.get_document",
        return_value={
            "document_id": "doc-123",
            "title": "Statement of Work",
            "file_type": "docx",
            "version": 2,
        },
    ):
        with TestClient(app) as client:
            resp = client.get("/api/documents/eagle/dev-tenant/packages/PKG-1/sow/v2/source.docx?content=true", headers=_auth_header())

    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is True
    assert "Existing scope paragraph" in data["content"]
    assert data["preview_mode"] == "docx_blocks"
    assert any(block["text"] == "Existing scope paragraph" for block in data["preview_blocks"])
    assert data["download_url"] == "https://signed.example/document"
    assert data["package_id"] == "PKG-1"
    assert data["document_type"] == "sow"
    assert data["document_id"] == "doc-123"
    assert data["version"] == 2
    assert data["file_type"] == "docx"


def test_put_document_rejects_binary_office_document():
    main_module = _load_app()
    app = main_module.app

    with TestClient(app) as client:
        resp = client.put(
            "/api/documents/eagle/dev-tenant/packages/PKG-1/sow/v2/source.docx",
            json={"content": "replacement text"},
            headers=_auth_header(),
        )

    assert resp.status_code == 415
    assert "plain text editor" in resp.json()["detail"].lower()


def test_get_document_still_returns_inline_content_for_text_documents():
    main_module = _load_app()
    app = main_module.app
    s3 = _build_s3_mock()

    with patch("boto3.client", return_value=s3):
        with TestClient(app) as client:
            resp = client.get("/api/documents/eagle/dev-tenant/dev-user/documents/test.md?content=true", headers=_auth_header())

    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is False
    assert data["content"].startswith("# Test markdown")
    assert data["download_url"] is None
    assert data["file_type"] == "md"


def test_get_document_returns_read_only_preview_for_xlsx():
    main_module = _load_app_without_docx_service()
    app = main_module.app
    s3 = _build_s3_mock()

    with patch("boto3.client", return_value=s3):
        with TestClient(app) as client:
            resp = client.get("/api/documents/eagle/dev-tenant/dev-user/documents/igce_20260316_120000.xlsx?content=true", headers=_auth_header())

    assert resp.status_code == 200
    data = resp.json()
    assert data["is_binary"] is True
    assert data["file_type"] == "xlsx"
    assert data["preview_mode"] == "xlsx_grid"
    assert "IGCE" in data["content"]
    assert "Olympus CK2" in data["content"]
    assert data["preview_sheets"]
    assert data["preview_sheets"][0]["title"] == "IGCE"
    assert data["download_url"] == "https://signed.example/document"


def test_post_docx_edit_updates_workspace_docx_preview_blocks():
    main_module = _load_app()
    app = main_module.app
    s3 = _build_s3_mock()
    import app.document_ai_edit_service as docx_service

    # Patch _get_s3 via the function's own globals (module reload makes
    # the canonical edit service module the correct patch target.
    _save_fn = docx_service.save_docx_preview_edits
    _orig_get_s3 = _save_fn.__globals__["_get_s3"]
    _save_fn.__globals__["_get_s3"] = lambda: s3

    try:
        with patch("boto3.client", return_value=s3), patch.dict(
            _save_fn.__globals__, {"write_document_changelog_entry": lambda **kw: None}
        ):
            headers = _auth_header()
            with TestClient(app) as client:
                get_resp = client.get("/api/documents/eagle/dev-tenant/dev-user/documents/test.docx?content=true", headers=headers)
                assert get_resp.status_code == 200
                blocks = get_resp.json()["preview_blocks"]
                paragraph_block = next(block for block in blocks if block["kind"] == "paragraph")
                paragraph_block["text"] = "Updated scope paragraph"

                save_resp = client.post(
                    "/api/documents/docx-edit/eagle/dev-tenant/dev-user/documents/test.docx",
                    json={"preview_blocks": blocks, "preview_mode": "docx_blocks"},
                    headers=headers,
                )
    finally:
        _save_fn.__globals__["_get_s3"] = _orig_get_s3

    assert save_resp.status_code == 200
    data = save_resp.json()
    assert data["success"] is True
    assert "Updated scope paragraph" in data["content"]
    assert any(block["text"] == "Updated scope paragraph" for block in data["preview_blocks"])


def test_post_xlsx_edit_updates_workspace_xlsx_preview_cells():
    main_module = _load_app()
    app = main_module.app
    s3 = _build_s3_mock()
    import app.spreadsheet_edit_service as xlsx_service

    # Patch _get_s3 via the function's own globals (module reload makes
    # the canonical spreadsheet service module the correct patch target.
    _save_fn = xlsx_service.save_xlsx_preview_edits
    _orig_get_s3 = _save_fn.__globals__["_get_s3"]
    _save_fn.__globals__["_get_s3"] = lambda: s3

    try:
        with patch("boto3.client", return_value=s3), patch.dict(
            _save_fn.__globals__, {"write_document_changelog_entry": lambda **kw: None}
        ):
            headers = _auth_header()
            with TestClient(app) as client:
                get_resp = client.get("/api/documents/eagle/dev-tenant/dev-user/documents/test.xlsx?content=true", headers=headers)
                assert get_resp.status_code == 200
                sheet = get_resp.json()["preview_sheets"][0]
                editable_cell = next(
                    cell
                    for row in sheet["rows"]
                    for cell in row["cells"]
                    if cell["editable"] and cell["cell_ref"] == "A2"
                )

                save_resp = client.post(
                    "/api/documents/xlsx-edit/eagle/dev-tenant/dev-user/documents/test.xlsx",
                    json={"cell_edits": [{"sheet_id": sheet["sheet_id"], "cell_ref": editable_cell["cell_ref"], "value": "Updated CK2"}]},
                    headers=headers,
                )
    finally:
        _save_fn.__globals__["_get_s3"] = _orig_get_s3

    assert save_resp.status_code == 200
    data = save_resp.json()
    assert data["success"] is True
    assert data["preview_mode"] == "xlsx_grid"
    assert any(
        cell["display_value"] == "Updated CK2"
        for row in data["preview_sheets"][0]["rows"]
        for cell in row["cells"]
        if cell["cell_ref"] == "A2"
    )
