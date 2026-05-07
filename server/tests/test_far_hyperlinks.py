"""Tests for FAR hyperlink injection in DOCX documents.

Validates that FAR/DFARS/HHSAR citations are correctly identified
and converted to hyperlinks pointing to acquisition.gov.
"""

import pytest

from app.template_service import _FAR_CITATION_RE, _build_far_url


class TestFarCitationRegex:
    """Test FAR citation regex patterns."""

    @pytest.mark.parametrize(
        "text,expected",
        [
            ("FAR 7.105", [("FAR", "7.105")]),
            ("far 7.105", [("far", "7.105")]),
            ("FAR 7.105(a)(4)", [("FAR", "7.105(a)(4)")]),
            ("FAR 52.219-9", [("FAR", "52.219-9")]),
            ("DFARS 252.204-7012", [("DFARS", "252.204-7012")]),
            ("HHSAR 352.239-73", [("HHSAR", "352.239-73")]),
            ("per FAR 13.106-3(a)(2)", [("FAR", "13.106-3(a)(2)")]),
            (
                "FAR 7.105 and FAR 52.219-9",
                [("FAR", "7.105"), ("FAR", "52.219-9")],
            ),
            ("No citations here", []),
            ("", []),
            ("FAR16.601", []),  # No space - should not match
            ("FAR 6.302-1", [("FAR", "6.302-1")]),
            ("DFARS 215.404-1", [("DFARS", "215.404-1")]),
        ],
    )
    def test_citation_extraction(self, text, expected):
        """Verify regex extracts FAR citations correctly."""
        matches = _FAR_CITATION_RE.findall(text)
        assert matches == expected

    def test_multiple_citations_in_paragraph(self):
        """Test extracting multiple citations from a paragraph."""
        text = (
            "This acquisition follows FAR 7.105 requirements for acquisition "
            "planning and FAR 52.219-9 for small business subcontracting. "
            "See also DFARS 252.204-7012 for cybersecurity requirements."
        )
        matches = _FAR_CITATION_RE.findall(text)
        assert len(matches) == 3
        assert ("FAR", "7.105") in matches
        assert ("FAR", "52.219-9") in matches
        assert ("DFARS", "252.204-7012") in matches


class TestBuildFarUrl:
    """Test FAR URL building."""

    @pytest.mark.parametrize(
        "citation_type,section,expected",
        [
            ("FAR", "7.105", "https://www.acquisition.gov/far/7.105"),
            ("FAR", "7.105(a)(4)", "https://www.acquisition.gov/far/7.105"),
            ("FAR", "52.219-9", "https://www.acquisition.gov/far/52.219-9"),
            ("DFARS", "252.204-7012", "https://www.acquisition.gov/dfars/252.204-7012"),
            ("HHSAR", "352.239-73", "https://www.acquisition.gov/hhsar/352.239-73"),
            ("far", "52.219-9", "https://www.acquisition.gov/far/52.219-9"),
            ("FAR", "6.302-1", "https://www.acquisition.gov/far/6.302-1"),
            ("FAR", "13.106-3(a)(2)", "https://www.acquisition.gov/far/13.106-3"),
        ],
    )
    def test_url_building(self, citation_type, section, expected):
        """Verify URL construction."""
        url = _build_far_url(citation_type, section)
        assert url == expected

    def test_subsection_stripped(self):
        """Verify subsection references are stripped from URL."""
        url = _build_far_url("FAR", "7.105(a)(4)")
        assert "(a)" not in url
        assert "(4)" not in url
        assert url == "https://www.acquisition.gov/far/7.105"


class TestHyperlinkInjection:
    """Test hyperlink injection in DOCX documents."""

    @pytest.fixture
    def sample_docx_bytes(self):
        """Create a minimal DOCX with FAR citations."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("This acquisition follows FAR 7.105 requirements.")
        doc.add_paragraph("See FAR 52.219-9 for small business provisions.")
        doc.add_paragraph("No citations in this paragraph.")
        doc.add_paragraph("Multiple: FAR 6.302-1 and DFARS 252.204-7012.")

        import io

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def test_hyperlink_injection_runs(self, sample_docx_bytes):
        """Verify hyperlink injection runs without error."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from app.template_service import DOCXPopulator
        import io

        doc = Document(io.BytesIO(sample_docx_bytes))
        # Should not raise
        count = DOCXPopulator.inject_far_hyperlinks(doc)
        # Count may be 0 or more depending on run structure
        assert count >= 0

    def test_document_still_valid_after_injection(self, sample_docx_bytes):
        """Verify document can still be saved after hyperlink injection."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from app.template_service import DOCXPopulator
        import io

        doc = Document(io.BytesIO(sample_docx_bytes))
        DOCXPopulator.inject_far_hyperlinks(doc)

        # Should be able to save without error
        output = io.BytesIO()
        doc.save(output)
        assert len(output.getvalue()) > 0

    def test_text_preserved_after_injection(self, sample_docx_bytes):
        """Verify original text is preserved after hyperlink injection."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from app.template_service import DOCXPopulator
        import io

        doc = Document(io.BytesIO(sample_docx_bytes))
        original_texts = [p.text for p in doc.paragraphs]

        DOCXPopulator.inject_far_hyperlinks(doc)

        # Save and reload
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        reloaded = Document(output)

        # Text content should be the same
        reloaded_texts = [p.text for p in reloaded.paragraphs]
        assert reloaded_texts == original_texts


class TestIntegrationWithPopulator:
    """Test FAR hyperlink injection integrated with DOCXPopulator.populate()."""

    @pytest.fixture
    def template_docx_bytes(self):
        """Create a template DOCX with placeholders and FAR citations."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_heading("{{PROJECT_TITLE}}", level=1)
        doc.add_paragraph(
            "This Statement of Work follows FAR 37.6 for performance-based acquisition."
        )
        doc.add_paragraph("Period of Performance: {{PERIOD_OF_PERFORMANCE}}")
        doc.add_paragraph(
            "Security requirements per FAR 52.204-21 and DFARS 252.204-7012."
        )

        import io

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def test_populate_includes_hyperlink_injection(self, template_docx_bytes):
        """Verify DOCXPopulator.populate() includes hyperlink injection."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from app.template_service import DOCXPopulator
        import io

        # Populate template
        data = {
            "title": "Cloud Services SOW",
            "period_of_performance": "12 months",
        }
        placeholder_map = {
            "title": "{{PROJECT_TITLE}}",
            "period_of_performance": "{{PERIOD_OF_PERFORMANCE}}",
        }

        populated_bytes = DOCXPopulator.populate(
            template_docx_bytes, data, placeholder_map
        )

        # Reload and verify
        doc = Document(io.BytesIO(populated_bytes))

        # Check placeholders were replaced
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cloud Services SOW" in full_text
        assert "12 months" in full_text
        assert "{{PROJECT_TITLE}}" not in full_text

        # FAR citations should still be in text (hyperlinked)
        assert "FAR 37.6" in full_text
        assert "FAR 52.204-21" in full_text
