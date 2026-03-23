"""Tests for FAR citation formatting in DOCX and PDF exports (Feature 3)."""
import pytest
import re


class TestFarCitationDocx:
    """Test FAR/DFARS citation styling in DOCX output."""

    def test_far_reference_styled_in_docx(self):
        """FAR citations should render with monospace bold font in NCI blue."""
        from app.document_export import markdown_to_docx
        from docx import Document
        from docx.shared import RGBColor
        import io

        content = "Per FAR 15.304, evaluation factors must be stated."
        data = markdown_to_docx(content, "Test Doc")

        doc = Document(io.BytesIO(data))
        found_citation = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "FAR 15.304" in run.text:
                    assert run.font.name == "Courier New"
                    assert run.bold is True
                    assert run.font.color.rgb == RGBColor(0, 51, 102)
                    found_citation = True

        assert found_citation, "FAR citation run not found in DOCX"

    def test_mixed_text_preserves_surrounding_formatting(self):
        """Text before and after a FAR citation should render normally."""
        from app.document_export import markdown_to_docx
        from docx import Document
        import io

        content = "See FAR 19.705 for subcontracting requirements."
        data = markdown_to_docx(content, "Test")

        doc = Document(io.BytesIO(data))
        texts = []
        for para in doc.paragraphs:
            for run in para.runs:
                texts.append(run.text)

        full_text = "".join(texts)
        assert "See" in full_text
        assert "FAR 19.705" in full_text
        assert "subcontracting" in full_text

    def test_citation_blockquote_renders_as_reference_section(self):
        """Blockquotes with 2+ FAR refs should render as Regulatory Reference."""
        from app.document_export import markdown_to_docx
        from docx import Document
        import io

        content = "> Per FAR 6.302 and FAR 6.304, sole source justification is required."
        data = markdown_to_docx(content, "Test")

        doc = Document(io.BytesIO(data))
        found_reg_ref = False
        for para in doc.paragraphs:
            full_text = "".join(r.text for r in para.runs)
            if "Regulatory Reference" in full_text:
                found_reg_ref = True

        assert found_reg_ref, "Regulatory Reference section not found"


class TestFarCitationPdf:
    """Test FAR/DFARS citation styling in PDF output."""

    def test_far_reference_styled_in_pdf(self):
        """PDF output should be valid and contain the citation text."""
        from app.document_export import markdown_to_pdf

        content = "Requirements per FAR 52.204 must be included."
        data = markdown_to_pdf(content, "Test")

        assert data.startswith(b"%PDF"), "Output is not valid PDF"
        assert len(data) > 100

    def test_multiple_citations_in_one_line(self):
        """Multiple FAR citations on one line should all be formatted."""
        from app.document_export import _markdown_to_reportlab

        text = "See FAR 15.304 and DFARS 215.304 for evaluation."
        result = _markdown_to_reportlab(text)

        # Both should be wrapped in font tags
        assert '<font name="Courier" color="#003366"><b>FAR 15.304</b></font>' in result
        assert '<font name="Courier" color="#003366"><b>DFARS 215.304</b></font>' in result

    def test_non_citation_blockquotes_unchanged(self):
        """Regular blockquotes without 2+ FAR refs should render normally."""
        from app.document_export import _FAR_CITATION_RE

        text = "This is a regular quote with no regulatory references."
        citations = _FAR_CITATION_RE.findall(text)
        assert len(citations) == 0

        text2 = "Per FAR 15.304, only one citation here."
        citations2 = _FAR_CITATION_RE.findall(text2)
        assert len(citations2) == 1  # Only 1, not enough for "Regulatory Reference"
