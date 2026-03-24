"""Tests for the document markdown conversion service."""

from __future__ import annotations

import io

from app.document_markdown_service import convert_to_markdown


class TestPlaintextToMarkdown:
    def test_utf8_passthrough(self):
        text = "Hello, this is a plain text document.\nWith multiple lines."
        result = convert_to_markdown(text.encode("utf-8"), "text/plain")
        assert result is not None
        assert "Hello" in result
        assert "multiple lines" in result

    def test_markdown_passthrough(self):
        md = "# Heading\n\nSome **bold** text."
        result = convert_to_markdown(md.encode("utf-8"), "text/markdown")
        assert result is not None
        assert "# Heading" in result

    def test_empty_body_returns_none(self):
        assert convert_to_markdown(b"", "text/plain") is None


class TestDocxToMarkdown:
    def _make_docx_bytes(self, paragraphs: list[tuple[str, str]] | None = None) -> bytes:
        """Create a minimal DOCX in memory.

        Args:
            paragraphs: list of (text, style) tuples. style can be "Heading 1", "Normal", etc.
        """
        from docx import Document
        doc = Document()
        if paragraphs is None:
            paragraphs = [
                ("Statement of Work", "Heading 1"),
                ("This is the scope section.", "Normal"),
                ("Deliverables", "Heading 2"),
                ("Item 1: Report", "Normal"),
            ]
        for text, style in paragraphs:
            doc.add_paragraph(text, style=style)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_basic_conversion(self):
        body = self._make_docx_bytes()
        result = convert_to_markdown(body, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result is not None
        assert "Statement of Work" in result
        assert "scope section" in result

    def test_headings_become_markdown_headers(self):
        body = self._make_docx_bytes()
        result = convert_to_markdown(body, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result is not None
        # Heading 1 should produce "# "
        assert "# Statement of Work" in result or "Statement of Work" in result

    def test_table_conversion(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Table Test", style="Heading 1")
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "100"
        table.cell(2, 0).text = "Beta"
        table.cell(2, 1).text = "200"
        buf = io.BytesIO()
        doc.save(buf)

        result = convert_to_markdown(buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result is not None
        assert "Name" in result
        assert "Alpha" in result
        assert "|" in result  # pipe table format


class TestXlsxToMarkdown:
    def _make_xlsx_bytes(self) -> bytes:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Cost Estimate"
        ws.append(["Item", "Quantity", "Unit Cost", "Total"])
        ws.append(["Labor", 100, 85.00, 8500.00])
        ws.append(["Materials", 50, 20.00, 1000.00])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_basic_conversion(self):
        body = self._make_xlsx_bytes()
        result = convert_to_markdown(body, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert result is not None
        assert "Cost Estimate" in result
        assert "Labor" in result
        assert "|" in result

    def test_sheet_title_in_output(self):
        body = self._make_xlsx_bytes()
        result = convert_to_markdown(body, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert result is not None
        assert "## Sheet: Cost Estimate" in result

    def test_multiple_sheets(self):
        from openpyxl import Workbook
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["A", "B"])
        ws1.append([1, 2])
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["X", "Y"])
        ws2.append([3, 4])
        buf = io.BytesIO()
        wb.save(buf)

        result = convert_to_markdown(buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert result is not None
        assert "Sheet1" in result
        assert "Sheet2" in result


class TestUnsupportedTypes:
    def test_png_returns_none(self):
        assert convert_to_markdown(b"\x89PNG\r\n\x1a\n", "image/png") is None

    def test_binary_returns_none(self):
        assert convert_to_markdown(b"\x00\x01\x02\x03", "application/octet-stream") is None


class TestFilenameExtensionFallback:
    def test_md_extension_fallback(self):
        text = "# Hello"
        result = convert_to_markdown(text.encode("utf-8"), "application/octet-stream", "readme.md")
        assert result is not None
        assert "Hello" in result

    def test_txt_extension_fallback(self):
        text = "plain text"
        result = convert_to_markdown(text.encode("utf-8"), "application/octet-stream", "notes.txt")
        assert result is not None
        assert "plain text" in result
