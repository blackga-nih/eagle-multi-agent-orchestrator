"""DOCX-aware AI document editing helpers."""

from __future__ import annotations

import io
import logging
import os
import re
import zipfile
from dataclasses import dataclass
from typing import Iterable, Optional

import boto3
from botocore.exceptions import BotoCoreError, ClientError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .changelog_store import write_document_changelog_entry
from .document_key_utils import (
    extract_package_document_ref,
    extract_workspace_document_ref,
    is_allowed_document_key,
)
from .document_service import create_package_document_version
from .document_store import get_document
from .template_service import DOCXPopulator

logger = logging.getLogger("eagle.document_ai_edit")

S3_BUCKET = os.getenv("S3_BUCKET", "eagle-documents-695681773636-dev")
AWS_REGION = os.getenv("AWS_REGION", "us-east-1")


@dataclass
class DocxEdit:
    search_text: str
    replacement_text: str


@dataclass
class DocxCheckboxEdit:
    label_text: str
    checked: bool


@dataclass
class DocxPreviewBlock:
    block_id: str
    kind: str
    text: str
    level: Optional[int] = None
    checked: Optional[bool] = None


def _get_s3():
    return boto3.client("s3", region_name=AWS_REGION)


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _extract_text_fallback(doc_bytes: bytes) -> Optional[str]:
    """Decode text-like bytes for misnamed .docx artifacts."""
    if not doc_bytes:
        return None

    try:
        decoded = doc_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return None

    stripped = decoded.strip()
    if not stripped:
        return None

    printable = sum(1 for ch in stripped if ch.isprintable() or ch in "\n\r\t")
    ratio = printable / max(len(stripped), 1)
    if ratio < 0.95:
        return None

    return stripped


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _is_truthy_word_value(value: Optional[str]) -> bool:
    return str(value or "").lower() in {"1", "true", "on"}


def _find_first_descendant_by_name(root, *names: str):
    wanted = set(names)
    for child in root.iter():
        if _local_name(child.tag) in wanted:
            return child
    return None


def _get_checkbox_state_from_paragraph(para) -> Optional[bool]:
    checkbox = _find_first_descendant_by_name(para._p, "checkBox", "checkbox")
    if checkbox is None:
        return None

    checked = _find_first_descendant_by_name(checkbox, "checked", "default")
    if checked is None:
        return False

    for key, value in checked.attrib.items():
        if _local_name(key) == "val":
            return _is_truthy_word_value(value)
    return False


def _set_checkbox_state_in_paragraph(para, checked: bool) -> bool:
    checkbox = _find_first_descendant_by_name(para._p, "checkBox", "checkbox")
    if checkbox is None:
        return False

    checkbox_name = _local_name(checkbox.tag)
    checked_node = _find_first_descendant_by_name(checkbox, "checked")
    attr_name = qn("w:val") if checkbox_name == "checkBox" else qn("w14:val")
    element_name = "w:checked" if checkbox_name == "checkBox" else "w14:checked"

    if checked_node is None:
        checked_node = OxmlElement(element_name)
        checkbox.append(checked_node)

    checked_node.set(attr_name, "1" if checked else "0")
    return True


def _paragraph_preview_text(para) -> Optional[str]:
    text = (para.text or "").strip()
    if not text and _get_checkbox_state_from_paragraph(para) is None:
        return None

    checkbox_state = _get_checkbox_state_from_paragraph(para)
    if checkbox_state is not None:
        mark = "x" if checkbox_state else " "
        return f"- [{mark}] {text}".rstrip()

    style_name = para.style.name if para.style else ""
    if "Heading 1" in style_name:
        return f"# {text}"
    if "Heading 2" in style_name:
        return f"## {text}"
    if "Heading 3" in style_name:
        return f"### {text}"
    return text or None


def _paragraph_block_kind(para) -> tuple[str, Optional[int], Optional[bool], str]:
    text = (para.text or "").strip()
    checkbox_state = _get_checkbox_state_from_paragraph(para)
    if checkbox_state is not None:
        return "checkbox", None, checkbox_state, text

    style_name = para.style.name if para.style else ""
    if "Heading 1" in style_name:
        return "heading", 1, None, text
    if "Heading 2" in style_name:
        return "heading", 2, None, text
    if "Heading 3" in style_name:
        return "heading", 3, None, text
    return "paragraph", None, None, text


def _serialize_preview_blocks(blocks: list[DocxPreviewBlock]) -> list[dict]:
    return [
        {
            "block_id": block.block_id,
            "kind": block.kind,
            "text": block.text,
            "level": block.level,
            "checked": block.checked,
        }
        for block in blocks
    ]


def _preview_text_from_blocks(blocks: list[DocxPreviewBlock]) -> str:
    lines: list[str] = []
    for block in blocks:
        if block.kind == "checkbox":
            mark = "x" if block.checked else " "
            lines.append(f"- [{mark}] {block.text}".rstrip())
        elif block.kind == "heading":
            level = max(1, min(block.level or 1, 6))
            lines.append(f"{'#' * level} {block.text}".rstrip())
        else:
            lines.append(block.text)
    return "\n\n".join(line for line in lines if line)


def _extract_docx_preview_blocks(docx_bytes: bytes) -> list[DocxPreviewBlock]:
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    blocks: list[DocxPreviewBlock] = []
    for idx, para in enumerate(_iter_document_paragraphs(doc)):
        kind, level, checked, text = _paragraph_block_kind(para)
        if not text and checked is None:
            continue
        blocks.append(
            DocxPreviewBlock(
                block_id=f"p:{idx}",
                kind=kind,
                text=text,
                level=level,
                checked=checked,
            )
        )
    return blocks


def _extract_text_preview_blocks(text: str) -> list[DocxPreviewBlock]:
    chunks = [chunk.strip() for chunk in re.split(r"\n\s*\n", text.strip()) if chunk.strip()]
    blocks: list[DocxPreviewBlock] = []
    for idx, chunk in enumerate(chunks):
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", chunk, re.DOTALL)
        checkbox_match = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)$", chunk, re.DOTALL)
        if heading_match:
            blocks.append(
                DocxPreviewBlock(
                    block_id=f"t:{idx}",
                    kind="heading",
                    text=heading_match.group(2).strip(),
                    level=len(heading_match.group(1)),
                )
            )
        elif checkbox_match:
            blocks.append(
                DocxPreviewBlock(
                    block_id=f"t:{idx}",
                    kind="checkbox",
                    text=checkbox_match.group(2).strip(),
                    checked=checkbox_match.group(1).lower() == "x",
                )
            )
        else:
            blocks.append(DocxPreviewBlock(block_id=f"t:{idx}", kind="paragraph", text=chunk))
    return blocks


def extract_docx_preview_payload(docx_bytes: bytes) -> dict:
    """Return preview text plus structured blocks for browser editing."""
    try:
        blocks = _extract_docx_preview_blocks(docx_bytes)
        return {
            "preview_mode": "docx_blocks",
            "content": _preview_text_from_blocks(blocks),
            "preview_blocks": _serialize_preview_blocks(blocks),
        }
    except zipfile.BadZipFile:
        logger.warning("DOCX preview fallback: artifact is not a valid zip package")
    except Exception:
        logger.exception("Failed to extract structured DOCX preview")

    fallback_text = _extract_text_fallback(docx_bytes)
    if fallback_text is None:
        return {"preview_mode": "none", "content": None, "preview_blocks": []}

    blocks = _extract_text_preview_blocks(fallback_text)
    return {
        "preview_mode": "text_fallback",
        "content": _preview_text_from_blocks(blocks),
        "preview_blocks": _serialize_preview_blocks(blocks),
    }


def extract_docx_preview(docx_bytes: bytes) -> Optional[str]:
    """Return a readable preview for a DOCX file."""
    payload = extract_docx_preview_payload(docx_bytes)
    return payload.get("content")


def _iter_document_paragraphs(doc) -> Iterable:
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    yield para
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    yield para


def _replace_paragraph_text(para, search_text: str, replacement_text: str) -> bool:
    current_text = para.text or ""
    if not current_text:
        return False

    if search_text in current_text:
        DOCXPopulator._replace_in_paragraph(para, {search_text: replacement_text})
        return True

    normalized_search = _normalize_text(search_text)
    normalized_current = _normalize_text(current_text)
    if normalized_search and normalized_search == normalized_current:
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run.text = ""
            first_run.text = replacement_text
        else:
            para.text = replacement_text
        return True

    return False


def _normalize_checkbox_label(text: str) -> str:
    cleaned = re.sub(r"^\s*[-*]?\s*\[[ xX]\]\s*", "", text or "")
    return _normalize_text(cleaned)


def _toggle_checkbox_by_label(para, label_text: str, checked: bool) -> bool:
    paragraph_label = _normalize_checkbox_label(para.text or "")
    target = _normalize_checkbox_label(label_text)
    if not paragraph_label or not target:
        return False
    if paragraph_label == target or target in paragraph_label:
        return _set_checkbox_state_in_paragraph(para, checked)
    return False


def _run_has_child(run, child_name: str) -> bool:
    return any(_local_name(child.tag) == child_name for child in run._r)


def _set_paragraph_visible_text(para, new_text: str) -> bool:
    text_runs = [run for run in para.runs if _run_has_child(run, "t")]
    if text_runs:
        text_runs[0].text = new_text
        for run in text_runs[1:]:
            run.text = ""
        return True

    if para.runs:
        para.runs[-1].text = new_text
        return True

    para.text = new_text
    return True


def apply_docx_edits(
    docx_bytes: bytes,
    edits: list[DocxEdit],
    checkbox_edits: Optional[list[DocxCheckboxEdit]] = None,
) -> tuple[bytes, int, list[str]]:
    """Apply targeted text replacements to a DOCX while preserving formatting where possible."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx required for DOCX editing") from exc

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = list(_iter_document_paragraphs(doc))
    missing: list[str] = []
    applied = 0

    for edit in edits:
        replaced = False
        for para in paragraphs:
            if _replace_paragraph_text(para, edit.search_text, edit.replacement_text):
                replaced = True
                applied += 1
                break
        if not replaced:
            missing.append(edit.search_text)

    for edit in checkbox_edits or []:
        toggled = False
        for para in paragraphs:
            if _toggle_checkbox_by_label(para, edit.label_text, edit.checked):
                toggled = True
                applied += 1
                break
        if not toggled:
            missing.append(f"checkbox:{edit.label_text}")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue(), applied, missing


def apply_docx_block_edits(docx_bytes: bytes, preview_blocks: list[dict], preview_mode: str) -> tuple[bytes, int]:
    """Apply browser-side block edits back to either a true DOCX or text fallback artifact."""
    if preview_mode == "text_fallback":
        blocks = [
            DocxPreviewBlock(
                block_id=str(block.get("block_id", "")),
                kind=str(block.get("kind", "paragraph")),
                text=str(block.get("text", "")),
                level=int(block["level"]) if block.get("level") is not None else None,
                checked=bool(block["checked"]) if block.get("checked") is not None else None,
            )
            for block in preview_blocks
        ]
        updated_text = _preview_text_from_blocks(blocks)
        return updated_text.encode("utf-8"), len(blocks)

    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = list(_iter_document_paragraphs(doc))
    applied = 0
    for block in preview_blocks:
        block_id = str(block.get("block_id", ""))
        if not block_id.startswith("p:"):
            continue
        try:
            para_idx = int(block_id.split(":", 1)[1])
        except (ValueError, IndexError):
            continue
        if para_idx < 0 or para_idx >= len(paragraphs):
            continue

        para = paragraphs[para_idx]
        kind, _level, current_checked, current_text = _paragraph_block_kind(para)
        next_text = str(block.get("text", ""))
        next_checked = block.get("checked")

        if kind == "checkbox" and isinstance(next_checked, bool) and next_checked != current_checked:
            if _set_checkbox_state_in_paragraph(para, next_checked):
                applied += 1

        if next_text != current_text:
            if _set_paragraph_visible_text(para, next_text):
                applied += 1

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue(), applied




def edit_docx_document(
    *,
    tenant_id: str,
    user_id: Optional[str],
    doc_key: str,
    edits: list[DocxEdit],
    checkbox_edits: Optional[list[DocxCheckboxEdit]] = None,
    session_id: Optional[str] = None,
    change_source: str = "ai_edit",
) -> dict:
    """Apply targeted edits to a DOCX document stored in S3."""
    if not doc_key:
        return {"error": "document_key is required"}
    if not edits and not checkbox_edits:
        return {"error": "At least one DOCX text edit or checkbox edit is required"}
    if not is_allowed_document_key(doc_key, tenant_id, user_id):
        return {"error": "Access denied for document key"}
    if not doc_key.lower().endswith(".docx"):
        return {"error": "DOCX edit tool only supports .docx documents"}

    s3 = _get_s3()
    try:
        response = s3.get_object(Bucket=S3_BUCKET, Key=doc_key)
        original_bytes = response["Body"].read()
    except (ClientError, BotoCoreError) as exc:
        logger.error("Failed to load DOCX for edit: %s", exc, exc_info=True)
        return {"error": f"Failed to load DOCX: {exc}"}

    try:
        updated_bytes, applied_count, missing = apply_docx_edits(original_bytes, edits, checkbox_edits)
    except Exception as exc:
        logger.error("Failed to apply DOCX edits: %s", exc, exc_info=True)
        return {"error": f"Failed to apply DOCX edits: {exc}"}

    if applied_count == 0:
        return {
            "error": "No DOCX edits were applied. Use exact existing text or checkbox labels from the current preview.",
            "missing": missing,
        }

    preview = extract_docx_preview(updated_bytes)
    package_ref = extract_package_document_ref(doc_key)
    if package_ref:
        if package_ref["tenant_id"] != tenant_id:
            return {"error": "Access denied for package document"}
        package_id = str(package_ref["package_id"])
        doc_type = str(package_ref["doc_type"])
        version = int(package_ref["version"])
        existing = get_document(tenant_id, package_id, doc_type, version)
        title = (existing or {}).get("title") or doc_type.replace("_", " ").title()

        result = create_package_document_version(
            tenant_id=tenant_id,
            package_id=package_id,
            doc_type=doc_type,
            content=updated_bytes,
            title=title,
            file_type="docx",
            created_by_user_id=user_id,
            session_id=session_id,
            change_source=change_source,
            template_id=(existing or {}).get("template_id"),
        )
        if not result.success:
            return {"error": result.error or "Failed to save DOCX edit"}

        return {
            "success": True,
            "mode": "package_docx_edit",
            "package_id": package_id,
            "document_id": result.document_id,
            "key": result.s3_key,
            "version": result.version,
            "file_type": "docx",
            "edits_applied": applied_count,
            "missing": missing,
            "content": preview,
            "message": f"Applied {applied_count} DOCX edit(s) and created version {result.version}.",
        }

    workspace_ref = extract_workspace_document_ref(doc_key)
    if not workspace_ref:
        return {"error": "Unsupported DOCX key format"}
    if workspace_ref["tenant"] != tenant_id or workspace_ref["user"] != user_id:
        return {"error": "Access denied for workspace document"}

    try:
        s3.put_object(
            Bucket=S3_BUCKET,
            Key=doc_key,
            Body=updated_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        write_document_changelog_entry(
            tenant_id=tenant_id,
            document_key=doc_key,
            change_type="update",
            change_source=change_source,
            change_summary=f"Applied {applied_count} AI DOCX edit(s)",
            actor_user_id=user_id or "ai-agent",
        )
    except Exception as exc:
        logger.error("Failed to persist workspace DOCX edit: %s", exc, exc_info=True)
        return {"error": f"Failed to save DOCX edit: {exc}"}

    return {
        "success": True,
        "mode": "workspace_docx_edit",
        "document_id": doc_key,
        "key": doc_key,
        "version": 0,
        "file_type": "docx",
        "edits_applied": applied_count,
        "missing": missing,
        "content": preview,
        "message": f"Applied {applied_count} DOCX edit(s).",
    }


def save_docx_preview_edits(
    *,
    tenant_id: str,
    user_id: Optional[str],
    doc_key: str,
    preview_blocks: list[dict],
    preview_mode: str,
    session_id: Optional[str] = None,
    change_source: str = "user_edit",
) -> dict:
    """Persist browser-side DOCX preview edits."""
    if not doc_key:
        return {"error": "document_key is required"}
    if not preview_blocks:
        return {"error": "preview_blocks are required"}
    if not is_allowed_document_key(doc_key, tenant_id, user_id):
        return {"error": "Access denied for document key"}
    if not doc_key.lower().endswith(".docx"):
        return {"error": "Structured preview editing only supports .docx documents"}

    s3 = _get_s3()
    try:
        response = s3.get_object(Bucket=S3_BUCKET, Key=doc_key)
        original_bytes = response["Body"].read()
        content_type = response.get("ContentType") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except (ClientError, BotoCoreError) as exc:
        logger.error("Failed to load DOCX preview artifact: %s", exc, exc_info=True)
        return {"error": "Failed to load document."}

    try:
        updated_bytes, applied_count = apply_docx_block_edits(original_bytes, preview_blocks, preview_mode)
    except Exception as exc:
        logger.error("Failed to apply structured DOCX preview edits: %s", exc, exc_info=True)
        return {"error": "Failed to apply preview edits."}

    if applied_count == 0:
        return {"error": "No preview edits were applied."}

    preview_payload = extract_docx_preview_payload(updated_bytes)
    package_ref = extract_package_document_ref(doc_key)
    if package_ref:
        if package_ref["tenant_id"] != tenant_id:
            return {"error": "Access denied for package document"}
        package_id = str(package_ref["package_id"])
        doc_type = str(package_ref["doc_type"])
        version = int(package_ref["version"])
        existing = get_document(tenant_id, package_id, doc_type, version)
        title = (existing or {}).get("title") or doc_type.replace("_", " ").title()
        result = create_package_document_version(
            tenant_id=tenant_id,
            package_id=package_id,
            doc_type=doc_type,
            content=updated_bytes,
            title=title,
            file_type="docx",
            created_by_user_id=user_id,
            session_id=session_id,
            change_source=change_source,
            template_id=(existing or {}).get("template_id"),
        )
        if not result.success:
            return {"error": result.error or "Failed to save document version"}
        return {
            "success": True,
            "mode": "package_docx_preview_edit",
            "document_id": result.document_id,
            "key": result.s3_key,
            "version": result.version,
            "file_type": "docx",
            "content": preview_payload.get("content"),
            "preview_blocks": preview_payload.get("preview_blocks", []),
            "preview_mode": preview_payload.get("preview_mode"),
            "message": f"Saved document version {result.version}.",
        }

    workspace_ref = extract_workspace_document_ref(doc_key)
    if not workspace_ref:
        return {"error": "Unsupported DOCX key format"}
    if workspace_ref["tenant"] != tenant_id or workspace_ref["user"] != user_id:
        return {"error": "Access denied for workspace document"}

    try:
        s3.put_object(
            Bucket=S3_BUCKET,
            Key=doc_key,
            Body=updated_bytes,
            ContentType=content_type,
        )
        write_document_changelog_entry(
            tenant_id=tenant_id,
            document_key=doc_key,
            change_type="update",
            change_source=change_source,
            change_summary="Updated DOCX via preview editor",
            actor_user_id=user_id or "user",
        )
    except Exception as exc:
        logger.error("Failed to save workspace DOCX preview edits: %s", exc, exc_info=True)
        return {"error": "Failed to save document."}

    return {
        "success": True,
        "mode": "workspace_docx_preview_edit",
        "document_id": doc_key,
        "key": doc_key,
        "version": 0,
        "file_type": "docx",
        "content": preview_payload.get("content"),
        "preview_blocks": preview_payload.get("preview_blocks", []),
        "preview_mode": preview_payload.get("preview_mode"),
        "message": "Document saved.",
    }
